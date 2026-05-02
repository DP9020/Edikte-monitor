"""
Erstellt brief_vorlage.docx aus dem Original-Brief (brief_vorlage_original.docx).
Ersetzt die personenbezogenen Inhalte durch {{PLATZHALTER}}.

Run: python3 create_brief_template.py
"""
import re
import sys

from docx import Document
from lxml import etree

doc = Document("brief_vorlage_original.docx")

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ── Inhalts-Validatoren für jeden hartcodierten Paragraph-Index ──────────────
# Wenn das Original-Word-Dokument umgeordnet wird, schlägt das Skript laut an
# (Index zeigt nicht mehr auf den erwarteten Inhalt) statt stillschweigend
# einen kaputten Brief zu erzeugen.
def _assert_paragraph(idx: int, contains_any: list[str]) -> None:
    para = doc.paragraphs[idx]
    text = para.text
    if not any(needle.lower() in text.lower() for needle in contains_any):
        raise SystemExit(
            f"❌ Paragraph[{idx}] enthält keines von {contains_any!r}. "
            f"Original-Brief vermutlich umgeordnet. Inhalt war: {text[:120]!r}"
        )


def replace_runs_text(para, new_text, bold=None):
    """Löscht alle Runs in einem Paragraphen, setzt einen neuen."""
    # Erst prüfen ob es eine Hyperlink-Struktur ist (keine Runs, aber w:hyperlink)
    ns = {"w": W}
    hyperlinks = para._element.findall(".//w:hyperlink", ns)

    if not para.runs and hyperlinks:
        # Hyperlink: alle w:t Elemente im Paragraphen ersetzen
        for t_el in para._element.findall(".//w:t", ns):
            t_el.text = new_text
            new_text = ""  # Rest leeren
        return

    if not para.runs:
        run = para.add_run(new_text)
        if bold is not None:
            run.bold = bold
        return
    first = para.runs[0]
    first.text = new_text
    if bold is not None:
        first.bold = bold
    for run in para.runs[1:]:
        run.text = ""

# ── Empfänger (paras 2–4) ────────────────────────────────────────────────────
replace_runs_text(doc.paragraphs[2], "{{EIGENTUEMER_NAME}}", bold=True)
replace_runs_text(doc.paragraphs[3], "{{ZUSTELL_ADRESSE}}", bold=True)
replace_runs_text(doc.paragraphs[4], "{{ZUSTELL_PLZ_ORT}}", bold=True)

# ── Datum (para 7, rechtsbündig) ─────────────────────────────────────────────
replace_runs_text(doc.paragraphs[7], "{{DATUM}}")

# ── Betreff: Liegenschaftsadresse (paras 10–11) ──────────────────────────────
replace_runs_text(doc.paragraphs[10], "{{LIEGENSCHAFT_ADRESSE}}", bold=True)
replace_runs_text(doc.paragraphs[11], "{{LIEGENSCHAFT_PLZ_ORT}}", bold=True)

# ── Anrede (para 12) ─────────────────────────────────────────────────────────
replace_runs_text(doc.paragraphs[12], "{{ANREDE}}")

# ── Fließtext para 33: Telefonnummer tauschen ────────────────────────────────
# Vor dem Replace prüfen, dass der Paragraph eine österreichische
# Telefonnummer enthält – sonst zeigt der Index woanders hin.
_assert_paragraph(33, ["+43"])
para33 = doc.paragraphs[33]
full = "".join(r.text for r in para33.runs)
new33 = re.sub(r'\+43\d[\d\s]+', "{{KONTAKT_TEL}}", full)
replace_runs_text(para33, new33)

# ── Signatur (paras 38–40) ───────────────────────────────────────────────────
replace_runs_text(doc.paragraphs[38], "{{KONTAKT_NAME}}")

# Para 39: Hyperlink – alle w:t Text-Elemente ersetzen
replace_runs_text(doc.paragraphs[39], "{{KONTAKT_EMAIL}}")

replace_runs_text(doc.paragraphs[40], "{{KONTAKT_TEL}}")

# ── Speichern ─────────────────────────────────────────────────────────────────
doc.save("brief_vorlage.docx")
print("✅ brief_vorlage.docx erstellt")

# ── Smoke-Test ────────────────────────────────────────────────────────────────
d = Document("brief_vorlage.docx")
all_text = "\n".join(p.text for p in d.paragraphs)

expected = [
    "{{EIGENTUEMER_NAME}}", "{{ZUSTELL_ADRESSE}}", "{{ZUSTELL_PLZ_ORT}}",
    "{{DATUM}}", "{{LIEGENSCHAFT_ADRESSE}}", "{{LIEGENSCHAFT_PLZ_ORT}}",
    "{{ANREDE}}", "{{KONTAKT_TEL}}", "{{KONTAKT_NAME}}",
]
missing: list[str] = []
for ph in expected:
    if ph in all_text:
        print(f"  ✅ {ph}")
    else:
        print(f"  ❌ FEHLT! {ph}")
        missing.append(ph)

# Email-Platzhalter ist in Hyperlink-XML, extra prüfen
xml_content = etree.tostring(d.element, encoding="unicode")
if "{{KONTAKT_EMAIL}}" in xml_content:
    print("  ✅ {{KONTAKT_EMAIL}} (Hyperlink)")
else:
    print("  ❌ FEHLT! {{KONTAKT_EMAIL}} (Hyperlink)")
    missing.append("{{KONTAKT_EMAIL}}")

if missing:
    sys.exit(f"❌ Template unvollständig – fehlende Platzhalter: {', '.join(missing)}")
