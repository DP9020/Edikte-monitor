"""
Edikte-Monitor – Österreich
============================
Scraper für https://edikte.justiz.gv.at (Gerichtliche Versteigerungen)
Alle Bundesländer | HTTP-Request (kein Browser nötig) | Notion | Telegram
"""

import os
import re
import json
import time
import asyncio
import base64
import urllib.request
import urllib.parse
from html import unescape as html_unescape
from datetime import datetime
from notion_client import Client

try:
    import fitz          # PyMuPDF – optionale Abhängigkeit
    FITZ_AVAILABLE = True
except ImportError:
    fitz = None
    FITZ_AVAILABLE = False

try:
    from openai import OpenAI as _OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    _OpenAI = None
    OPENAI_AVAILABLE = False

try:
    from googleapiclient.discovery import build as _gdrive_build
    from googleapiclient.http import MediaIoBaseUpload
    from google.oauth2 import service_account as _gsa
    import io as _io
    GDRIVE_AVAILABLE = True
except ImportError:
    GDRIVE_AVAILABLE = False

# =============================================================================
# KONFIGURATION
# =============================================================================

BASE_URL = "https://edikte.justiz.gv.at"

# Bundesland-Werte aus dem Formular (name=BL)
BUNDESLAENDER = {
    "Wien":           "0",
    "Niederösterreich": "1",
    "Burgenland":     "2",
    "Oberösterreich": "3",
    "Salzburg":       "4",
    "Steiermark":     "5",
    "Kärnten":        "6",
    "Tirol":          "7",
    "Vorarlberg":     "8",
}

# Nur diese Link-Texte werden verarbeitet
RELEVANT_TYPES = ("Versteigerung", "Entfall des Termins", "Verschiebung")

# Schlüsselwörter im Link-Text → Objekt wird NICHT importiert
# (greift auf Ergebnisseite, wo der Text oft nur "Versteigerung (Datum)" ist)
EXCLUDE_KEYWORDS = [
    "landwirtschaft",
    "land- und forst",
    "forstwirtschaft",
    "gewerb",
    "betriebsobjekt",
    "industrie",
    "lager",
    "büro",
    "hotel",
    "pension",
]

# Kategorien aus der Detailseite → Objekt wird NICHT importiert
# Entspricht den Werten im Feld "Kategorie(n)" auf edikte.justiz.gv.at
EXCLUDE_KATEGORIEN = {
    "land- und forstwirtschaftlich genutzte liegenschaft",  # LF
    "gewerbliche liegenschaft",                             # GL
    "betriebsobjekt",
    "superädifikat",                                        # SE – nur wenn gewerblich
}

# Notion-Feldname für PLZ (exakt so wie in der Datenbank angelegt)
NOTION_PLZ_FIELD = "Liegenschafts PLZ"

# Workflow-Phasen die NICHT automatisch überschrieben werden dürfen
# (Einträge die bereits manuell bearbeitet wurden)
GESCHUETZT_PHASEN: frozenset[str] = frozenset({
    "🔎 In Prüfung",
    "❌ Nicht relevant",
    "✅ Relevant – Brief vorbereiten",
    "📩 Brief versendet",
    "📊 Gutachten analysiert",
    "✅ Gekauft",
    "🗄 Archiviert",
})

# Edikt-ID aus dem Link extrahieren
ID_RE = re.compile(r"alldoc/([0-9a-f]+)!OpenDocument", re.IGNORECASE)

# Verkehrswert / Schätzwert
SCHAETZWERT_RE = re.compile(
    r'(?:Schätzwert|Verkehrswert|Schätzungswert|Wert)[:\s]+([\d\.\s,]+(?:EUR|€)?)',
    re.IGNORECASE
)


# =============================================================================
# HILFSFUNKTIONEN
# =============================================================================

def env(name: str) -> str:
    """Liest eine Umgebungsvariable – wirft Fehler wenn nicht gesetzt."""
    value = os.getenv(name)
    if not value:
        raise RuntimeError(f"Fehlende Umgebungsvariable: {name}")
    return value


def clean_notion_db_id(raw: str) -> str:
    """Bereinigt die Notion Datenbank-ID (entfernt View-Parameter etc.)."""
    raw = raw.split("?")[0].strip()
    raw = raw.rstrip("/").split("/")[-1]
    clean = re.sub(r"[^0-9a-fA-F]", "", raw)
    if len(clean) == 32:
        return f"{clean[0:8]}-{clean[8:12]}-{clean[12:16]}-{clean[16:20]}-{clean[20:32]}"
    return raw


def is_excluded(text: str) -> bool:
    """Prüft ob ein Objekt anhand des Link-Texts ausgeschlossen werden soll."""
    return any(kw in text.lower() for kw in EXCLUDE_KEYWORDS)


def is_excluded_by_kategorie(kategorie: str) -> bool:
    """Prüft ob ein Objekt anhand der Detailseiten-Kategorie ausgeschlossen werden soll."""
    return kategorie.lower().strip() in EXCLUDE_KATEGORIEN


def parse_euro(raw: str) -> float | None:
    """
    Wandelt einen österreichischen Betragsstring in float um.
    z.B. '180.000,00 EUR' → 180000.0
    """
    try:
        cleaned = re.sub(r"[€EUReur\s]", "", raw.strip())
        cleaned = cleaned.replace(".", "").replace(",", ".")
        return float(cleaned)
    except Exception:
        return None


def parse_flaeche(raw: str) -> float | None:
    """Wandelt '96,72 m²' in 96.72 um."""
    try:
        m = re.search(r"([\d.,]+)", raw)
        if m:
            return float(m.group(1).replace(".", "").replace(",", "."))
    except Exception:
        pass
    return None


def fetch_detail(link: str) -> dict:
    """
    Lädt die Edikt-Detailseite und extrahiert alle strukturierten Felder
    direkt aus dem Bootstrap-Grid (span.col-sm-3 + p.col-sm-9).

    Liefert ein Dict mit den Schlüsseln:
      liegenschaftsadresse, plz_ort, adresse_voll   ← echte Immobilienadresse
      gericht, aktenzeichen, wegen
      termin, termin_iso
      kategorie, grundbuch, ez
      flaeche_objekt, flaeche_grundstueck
      schaetzwert (float), schaetzwert_str
      geringstes_gebot (float)
    """
    try:
        req = urllib.request.Request(
            link,
            headers={"User-Agent": "Mozilla/5.0 (compatible; EdikteMonitor/1.0)"}
        )
        with urllib.request.urlopen(req, timeout=20) as r:
            html = r.read().decode("utf-8", errors="replace")
    except Exception as exc:
        print(f"    [Detail] ⚠️  Fehler beim Laden: {exc}")
        return {}

    # ── Alle label→value Paare aus dem Bootstrap-Grid extrahieren ────────────
    grid_re = re.compile(
        r'<span[^>]*col-sm-3[^>]*>\s*([^<]+?)\s*</span>\s*<p[^>]*col-sm-9[^>]*>\s*(.*?)\s*</p>',
        re.DOTALL | re.IGNORECASE
    )

    def clean(html_fragment: str) -> str:
        t = re.sub(r"<[^>]+>", " ", html_fragment)
        t = t.replace("\xa0", " ").replace("&nbsp;", " ")
        t = html_unescape(t)
        return " ".join(t.split()).strip()

    fields: dict[str, str] = {}
    for label, value in grid_re.findall(html):
        key = label.strip().rstrip(":").strip()
        fields[key] = clean(value)

    result: dict = {}

    # ── Liegenschaftsadresse (echte Immobilienadresse!) ──────────────────────
    adresse    = fields.get("Liegenschaftsadresse", "")
    plz_ort    = fields.get("PLZ/Ort", "")
    if adresse:
        result["liegenschaftsadresse"] = adresse
        result["plz_ort"]              = plz_ort
        result["adresse_voll"]         = f"{adresse}, {plz_ort}".strip(", ")
        print(f"    [Detail] 📍 Adresse: {result['adresse_voll']}")

    # ── Gericht / Dienststelle ────────────────────────────────────────────────
    if "Dienststelle" in fields:
        result["gericht"] = fields["Dienststelle"]
    elif "Dienststelle:" in fields:
        result["gericht"] = fields["Dienststelle:"]

    # ── Aktenzeichen ──────────────────────────────────────────────────────────
    for k in ("Aktenzeichen", "Aktenzeichen:"):
        if k in fields:
            result["aktenzeichen"] = fields[k]
            break

    # ── wegen ─────────────────────────────────────────────────────────────────
    if "wegen" in fields:
        result["wegen"] = fields["wegen"]

    # ── Versteigerungstermin ──────────────────────────────────────────────────
    termin_raw = fields.get("Versteigerungstermin", "")
    m = re.search(r"(\d{1,2}\.\d{1,2}\.\d{4})\s+um\s+([\d:]+\s*Uhr)", termin_raw)
    if m:
        result["termin"] = f"{m.group(1)} {m.group(2)}"
        try:
            dt = datetime.strptime(m.group(1), "%d.%m.%Y")
            result["termin_iso"] = dt.strftime("%Y-%m-%d")
        except Exception:
            pass

    # ── Kategorie / Objektart ─────────────────────────────────────────────────
    if "Kategorie(n)" in fields:
        result["kategorie"] = fields["Kategorie(n)"]

    # ── Grundbuch / EZ ────────────────────────────────────────────────────────
    if "Grundbuch" in fields:
        result["grundbuch"] = fields["Grundbuch"]
    if "EZ" in fields:
        result["ez"] = fields["EZ"]

    # ── Flächen ───────────────────────────────────────────────────────────────
    fobj = fields.get("Objektgröße", "")
    if fobj:
        parsed = parse_flaeche(fobj)
        if parsed:
            result["flaeche_objekt"] = parsed

    fgrst = fields.get("Grundstücksgröße", "")
    if fgrst:
        parsed = parse_flaeche(fgrst)
        if parsed:
            result["flaeche_grundstueck"] = parsed

    # ── Schätzwert ────────────────────────────────────────────────────────────
    sv_raw = fields.get("Schätzwert", "")
    if sv_raw:
        result["schaetzwert_str"] = sv_raw
        parsed = parse_euro(sv_raw)
        if parsed is not None:
            result["schaetzwert"] = parsed
            print(f"    [Detail] 💰 Schätzwert: {parsed:,.0f} €")

    # ── Geringstes Gebot ──────────────────────────────────────────────────────
    gg_raw = fields.get("Geringstes Gebot", "")
    if gg_raw:
        parsed = parse_euro(gg_raw)
        if parsed is not None:
            result["geringstes_gebot"] = parsed

    return result


# =============================================================================
# GLOBALE HILFSFUNKTIONEN (werden von mehreren Modulen genutzt)
# =============================================================================

def _rt(text: str) -> dict:
    """Erstellt ein Notion rich_text Property."""
    return {"rich_text": [{"text": {"content": str(text)[:2000]}}]}


def _str_val(val) -> str:
    """Konvertiert einen Wert sicher zu str."""
    return str(val).strip() if val else ""


def _lst_val(val) -> list:
    """Konvertiert einen Wert sicher zu einer bereinigten Liste."""
    if isinstance(val, list):
        return [str(v).strip() for v in val if v and str(v).strip()]
    if isinstance(val, str) and val.strip():
        return [val.strip()]
    return []


def _clean_name(name: str) -> str:
    """Verwirft Parser-Artefakte die als Eigentümername durchgerutscht sind."""
    if not name:
        return ""
    INVALID = {"nicht angegeben", "unbekannt", "n/a", "none", "null", "-", "–"}
    if name.strip().lower() in INVALID:
        return ""
    if re.match(r'^[)\]}>]', name) or name.rstrip().endswith('-'):
        return ""
    if not any(c.isalpha() for c in name):
        return ""
    return name


def _clean_adresse(adr: str) -> str:
    """Bereinigt fehlerhafte Adressen aus der PDF-Extraktion."""
    if not adr:
        return ""
    adr = re.sub(r',?\s*Telefon.*$', '', adr, flags=re.IGNORECASE).strip().rstrip(',')
    m = re.match(r'^(?:[A-Za-z]-?)?\d{4,5}\s+\S+.*?,\s*(.+)', adr)
    if m:
        adr = m.group(1).strip()
    adr = re.sub(r',\s*[A-ZÄÖÜ][a-zäöüß]+$', '', adr).strip()
    return adr


# =============================================================================
# TELEGRAM
# =============================================================================

def html_escape(text: str) -> str:
    """Escapt Sonderzeichen für Telegram HTML-Modus."""
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;"))


def _telegram_send_raw(url: str, payload_dict: dict) -> None:
    """Interne Hilfsfunktion: sendet einen JSON-Payload an die Telegram API."""
    payload = json.dumps(payload_dict, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=payload,
        headers={"Content-Type": "application/json; charset=utf-8"},
    )
    with urllib.request.urlopen(req, timeout=15) as r:
        r.read()


def _truncate_plain(text: str, limit: int = 4096) -> str:
    """Kürzt Plain-Text sicher auf das Zeichenlimit."""
    if len(text) <= limit:
        return text
    return text[:limit - 6] + "\n[...]"


def _strip_html_tags(text: str) -> str:
    """Entfernt alle HTML-Tags und dekodiert HTML-Entities."""
    plain = re.sub(r"<[^>]+>", "", text)
    plain = plain.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    return plain


# Bundesländer die Benjamin (Pippan) betreffen
BENJAMIN_BUNDESLAENDER = {"Wien", "Oberösterreich"}

# Bundesländer die Christopher (Dovjak) betreffen
CHRISTOPHER_BUNDESLAENDER = {"Niederösterreich", "Burgenland"}


def _get_benjamin_chat_id() -> str:
    """Gibt die Telegram Chat-ID von Benjamin zurück (aus Umgebungsvariable)."""
    return os.environ.get("TELEGRAM_CHAT_ID_BENJAMIN", "")


def _get_christopher_chat_id() -> str:
    """Gibt die Telegram Chat-ID von Christopher zurück (aus Umgebungsvariable)."""
    return os.environ.get("TELEGRAM_CHAT_ID_CHRISTOPHER", "")


def send_telegram_document(docx_bytes: bytes, dateiname: str, caption: str = "", bundesland: str = "") -> bool:
    """
    Schickt eine DOCX-Datei als Telegram-Dokument (sendDocument, multipart/form-data).
    Wenn bundesland in BENJAMIN_BUNDESLAENDER → auch an Benjamin senden.
    Gibt True zurück wenn erfolgreich, sonst False.
    """
    try:
        token   = env("TELEGRAM_BOT_TOKEN")
        chat_id = env("TELEGRAM_CHAT_ID")
        url     = f"https://api.telegram.org/bot{token}/sendDocument"

        boundary = "----TelegramBoundary7438291"
        CRLF = b"\r\n"

        def field(name: str, value: str) -> bytes:
            return (
                f"--{boundary}\r\n"
                f'Content-Disposition: form-data; name="{name}"\r\n\r\n'
                f"{value}\r\n"
            ).encode("utf-8")

        body = (
            field("chat_id", chat_id)
            + field("caption", caption[:1024])
            + (
                f"--{boundary}\r\n"
                f'Content-Disposition: form-data; name="document"; filename="{dateiname}"\r\n'
                f"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
            ).encode("utf-8")
            + docx_bytes
            + CRLF
            + f"--{boundary}--\r\n".encode("utf-8")
        )

        req = urllib.request.Request(
            url,
            data=body,
            headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        )
        with urllib.request.urlopen(req, timeout=30) as r:
            r.read()
        print(f"  [Brief] 📨 Telegram-Dokument gesendet: {dateiname}")

        # Auch an Benjamin senden wenn Bundesland Wien oder OÖ
        benjamin_id = _get_benjamin_chat_id()
        if benjamin_id and bundesland in BENJAMIN_BUNDESLAENDER:
            _send_document_to_chat(token, benjamin_id, docx_bytes, dateiname, caption)
            print(f"  [Brief] 📨 Telegram-Dokument auch an Benjamin gesendet: {dateiname}")

        # Auch an Christopher senden wenn Bundesland NÖ oder Burgenland
        christopher_id = _get_christopher_chat_id()
        if christopher_id and bundesland in CHRISTOPHER_BUNDESLAENDER:
            _send_document_to_chat(token, christopher_id, docx_bytes, dateiname, caption)
            print(f"  [Brief] 📨 Telegram-Dokument auch an Christopher gesendet: {dateiname}")

        return True
    except Exception as exc:
        print(f"  [Brief] ⚠️  Telegram-Dokument fehlgeschlagen: {exc}")
        return False


def _send_document_to_chat(token: str, chat_id: str, docx_bytes: bytes, dateiname: str, caption: str) -> None:
    """Sendet ein Dokument an eine bestimmte Chat-ID."""
    url      = f"https://api.telegram.org/bot{token}/sendDocument"
    boundary = "----TelegramBoundary7438292"
    CRLF     = b"\r\n"

    def field(name: str, value: str) -> bytes:
        return (
            f"--{boundary}\r\n"
            f'Content-Disposition: form-data; name="{name}"\r\n\r\n'
            f"{value}\r\n"
        ).encode("utf-8")

    body = (
        field("chat_id", chat_id)
        + field("caption", caption[:1024])
        + (
            f"--{boundary}\r\n"
            f'Content-Disposition: form-data; name="document"; filename="{dateiname}"\r\n'
            f"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
        ).encode("utf-8")
        + docx_bytes
        + CRLF
        + f"--{boundary}--\r\n".encode("utf-8")
    )
    req = urllib.request.Request(
        url, data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
    )
    with urllib.request.urlopen(req, timeout=30) as r:
        r.read()


async def send_telegram(message: str, extra_chat_ids: list[str] | None = None) -> None:
    """
    Sendet eine Nachricht via Telegram Bot (HTML-Modus).
    - Wenn die Nachricht > 4096 Zeichen: wird in mehrere Teile aufgeteilt.
    - Bei HTML-Fehler (400): Fallback auf reinen Text ohne parse_mode.
    - extra_chat_ids: zusätzliche Chat-IDs die dieselbe Nachricht bekommen.
    """
    token   = env("TELEGRAM_BOT_TOKEN")
    chat_id = env("TELEGRAM_CHAT_ID")
    url     = f"https://api.telegram.org/bot{token}/sendMessage"

    def split_message(text: str, limit: int = 4000) -> list[str]:
        """Teilt eine Nachricht an Zeilengrenzen auf, sodass kein HTML-Tag zerrissen wird."""
        if len(text) <= limit:
            return [text]
        parts = []
        current = []
        current_len = 0
        for line in text.split("\n"):
            line_len = len(line) + 1  # +1 für \n
            if current_len + line_len > limit and current:
                parts.append("\n".join(current))
                current = [line]
                current_len = line_len
            else:
                current.append(line)
                current_len += line_len
        if current:
            parts.append("\n".join(current))
        return parts

    parts = split_message(message)
    total = len(parts)

    for i, part in enumerate(parts, 1):
        label = f" ({i}/{total})" if total > 1 else ""
        try:
            _telegram_send_raw(url, {
                "chat_id":                  chat_id,
                "text":                     part,
                "parse_mode":               "HTML",
                "disable_web_page_preview": True,
            })
            print(f"[Telegram] ✅ Nachricht{label} gesendet ({len(part)} Zeichen)")
        except Exception as e:
            print(f"[Telegram] ⚠️  HTML-Modus fehlgeschlagen{label} ({e}), versuche Plain Text …")
            # Fallback: HTML-Tags entfernen, kein parse_mode senden
            plain = _truncate_plain(_strip_html_tags(part))
            try:
                _telegram_send_raw(url, {
                    "chat_id":                  chat_id,
                    "text":                     plain,
                    "disable_web_page_preview": True,
                })
                print(f"[Telegram] ✅ Plain-Text{label} gesendet ({len(plain)} Zeichen)")
            except Exception as e2:
                raise RuntimeError(f"Telegram komplett fehlgeschlagen{label}: {e2}") from e2

    # Nachricht auch an extra Chat-IDs senden (z.B. Benjamin)
    if extra_chat_ids:
        for extra_id in extra_chat_ids:
            for i, part in enumerate(parts, 1):
                label = f" ({i}/{total})" if total > 1 else ""
                try:
                    _telegram_send_raw(url, {
                        "chat_id":                  extra_id,
                        "text":                     part,
                        "parse_mode":               "HTML",
                        "disable_web_page_preview": True,
                    })
                    print(f"[Telegram] ✅ Nachricht{label} an {extra_id} gesendet")
                except Exception as e:
                    plain = _truncate_plain(_strip_html_tags(part))
                    try:
                        _telegram_send_raw(url, {
                            "chat_id":                  extra_id,
                            "text":                     plain,
                            "disable_web_page_preview": True,
                        })
                    except Exception:
                        print(f"[Telegram] ⚠️  Nachricht an {extra_id} fehlgeschlagen: {e}")


# =============================================================================
# GOOGLE DRIVE – Unterlagen-Upload für "Gelb"-Einträge
# =============================================================================

def gdrive_get_service():
    """Erstellt den Google Drive API Service via Service Account (Base64 oder raw JSON)."""
    print(f"[GDrive] 🔧 GDRIVE_AVAILABLE={GDRIVE_AVAILABLE}")
    if not GDRIVE_AVAILABLE:
        print("[GDrive] ❌ google-api-python-client nicht installiert")
        return None
    key_raw = os.environ.get("GOOGLE_SERVICE_ACCOUNT_KEY", "")
    if not key_raw:
        print("[GDrive] ❌ GOOGLE_SERVICE_ACCOUNT_KEY nicht gesetzt")
        return None
    print(f"[GDrive] 🔑 Key gefunden (Länge: {len(key_raw)} Zeichen, Base64={not key_raw.strip().startswith('{')})")
    folder_id = os.environ.get("GOOGLE_DRIVE_FOLDER_ID", "")
    print(f"[GDrive] 📁 GOOGLE_DRIVE_FOLDER_ID={'gesetzt ('+folder_id[:8]+'…)' if folder_id else 'NICHT GESETZT'}")
    try:
        if not key_raw.strip().startswith("{"):
            # Fehlende Padding-Zeichen ergänzen (= am Ende)
            padded = key_raw.strip() + "=" * (4 - len(key_raw.strip()) % 4)
            key_raw = base64.b64decode(padded).decode("utf-8")
        creds_info = json.loads(key_raw)
        print(f"[GDrive] 👤 Service Account: {creds_info.get('client_email', '?')}")
        creds = _gsa.Credentials.from_service_account_info(
            creds_info,
            scopes=["https://www.googleapis.com/auth/drive"],
        )
        svc = _gdrive_build("drive", "v3", credentials=creds, cache_discovery=False)
        print("[GDrive] ✅ Service erfolgreich erstellt")
        return svc
    except Exception as exc:
        print(f"[GDrive] ❌ Service-Erstellung fehlgeschlagen: {exc}")
        return None


def gdrive_find_or_create_folder(service, name: str, parent_id: str) -> str:
    """Gibt die ID eines Google-Drive-Ordners zurück – erstellt ihn falls nötig."""
    safe_name = name.replace("'", "\\'")
    query = (
        f"name='{safe_name}' "
        f"and mimeType='application/vnd.google-apps.folder' "
        f"and '{parent_id}' in parents "
        f"and trashed=false"
    )
    result = service.files().list(q=query, fields="files(id)", pageSize=1).execute()
    files  = result.get("files", [])
    if files:
        return files[0]["id"]
    meta   = {
        "name":     name,
        "mimeType": "application/vnd.google-apps.folder",
        "parents":  [parent_id],
    }
    folder = service.files().create(body=meta, fields="id").execute()
    return folder["id"]


def gdrive_upload_file(service, data: bytes, filename: str, folder_id: str) -> str:
    """Lädt Bytes als Datei in Google Drive hoch, gibt File-ID zurück."""
    ext      = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    mime_map = {
        "pdf":  "application/pdf",
        "jpg":  "image/jpeg",
        "jpeg": "image/jpeg",
        "png":  "image/png",
    }
    mime  = mime_map.get(ext, "application/octet-stream")
    meta  = {"name": filename, "parents": [folder_id]}
    media = MediaIoBaseUpload(_io.BytesIO(data), mimetype=mime, resumable=False)
    f     = service.files().create(body=meta, media_body=media, fields="id").execute()
    return f["id"]


def gdrive_clear_placeholder_links(notion: Client, db_id: str, all_pages: list[dict]) -> int:
    """
    Entfernt fälschlicherweise gesetzte Platzhalter-Links
    ('nicht-verfuegbar') aus dem Google Drive Link Feld.
    Wird einmalig nach einem fehlgeschlagenen Run benötigt.
    """
    PLACEHOLDER = "https://drive.google.com/drive/folders/nicht-verfuegbar"
    cleared = 0
    for page in all_pages:
        props      = page.get("properties", {})
        drive_link = props.get("Google Drive Link", {}).get("url") or ""
        status     = (props.get("Status", {}).get("select") or {}).get("name", "")
        if drive_link == PLACEHOLDER and status == "🟡 Gelb":
            try:
                notion.pages.update(
                    page_id=page["id"],
                    properties={"Google Drive Link": {"url": None}},
                )
                cleared += 1
            except Exception as exc:
                print(f"  [GDrive] ⚠️  Platzhalter-Bereinigung fehlgeschlagen: {exc}")
    if cleared:
        print(f"[GDrive] 🧹 {cleared} Platzhalter-Links bereinigt")
    return cleared


def gdrive_sync_gelb_entries(
    notion: Client, db_id: str, all_pages: list[dict], service
) -> int:
    """
    Lädt für alle Einträge mit Status '🟡 Gelb' (und noch keinem Drive-Link)
    alle Unterlagen von der Edikt-Seite in den konfigurierten Google-Drive-Ordner hoch.
    Speichert den Ordner-Link danach im Notion-Feld 'Google Drive Link'.
    Gibt die Anzahl erfolgreich bearbeiteter Einträge zurück.
    """
    parent_folder_id = os.environ.get("GOOGLE_DRIVE_FOLDER_ID", "")
    if not parent_folder_id:
        print("[GDrive] ℹ️  GOOGLE_DRIVE_FOLDER_ID nicht gesetzt – überspringe.")
        return 0

    kandidaten: list[dict] = []
    gelb_gesamt = 0
    for page in all_pages:
        props       = page.get("properties", {})
        status      = (props.get("Status", {}).get("select") or {}).get("name", "")
        edikt_link  = props.get("Link", {}).get("url") or ""
        drive_link  = props.get("Google Drive Link", {}).get("url") or ""
        if status == "🟡 Gelb":
            gelb_gesamt += 1
            if not edikt_link:
                print(f"  [GDrive] ⚠️  Gelb-Eintrag ohne Link übersprungen")
            elif drive_link:
                print(f"  [GDrive] ℹ️  Drive-Link bereits vorhanden – übersprungen")
            else:
                kandidaten.append(page)

    print(f"\n[GDrive] 🔍 {gelb_gesamt} Gelb-Einträge gesamt, {len(kandidaten)} ohne Drive-Link")
    if not kandidaten:
        return 0

    erledigt = 0
    for page in kandidaten:
        props = page.get("properties", {})

        # Adresse (Title-Feld)
        title_rt = props.get("Liegenschaftsadresse", {}).get("title", [])
        adresse  = title_rt[0].get("plain_text", "").strip() if title_rt else ""

        # Verpflichtende Partei (Eigentümer)
        vp_rt    = props.get("Verpflichtende Partei", {}).get("rich_text", [])
        vp_name  = "".join(t.get("text", {}).get("content", "") for t in vp_rt).strip()

        edikt_url = props.get("Link", {}).get("url") or ""
        page_id   = page["id"]

        # Ordnernamen erstellen und für Drive bereinigen
        raw_name    = f"{vp_name} - {adresse}" if vp_name else adresse
        folder_name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", raw_name)[:200].strip(" .-")
        if not folder_name:
            folder_name = page_id[:12]

        print(f"\n[GDrive] 📁 Verarbeite: {folder_name}")

        # ── Schritt 1: Anhänge von Edikt-Seite holen (vor Drive-Zugriff) ──────
        # Fehler hier = Edikt-Seite weg → Platzhalter setzen
        try:
            attachments = gutachten_fetch_attachment_links(edikt_url)
            all_files   = attachments.get("pdfs", []) + attachments.get("images", [])
            print(f"  [GDrive] 📎 {len(all_files)} Datei(en) auf Edikt-Seite gefunden")
        except Exception as fetch_exc:
            print(f"  [GDrive] ⚠️  Edikt-Seite nicht erreichbar: {fetch_exc}")
            try:
                notion.pages.update(
                    page_id=page_id,
                    properties={"Google Drive Link": {"url": "https://drive.google.com/drive/folders/nicht-verfuegbar"}},
                )
                print(f"  [GDrive] ℹ️  Platzhalter-Link gesetzt (Edikt-Seite nicht erreichbar)")
            except Exception:
                pass
            time.sleep(0.5)
            continue

        # ── Schritt 2: Drive-Ordner anlegen + Dateien hochladen ───────────────
        # Fehler hier = Drive-API-Problem → KEIN Platzhalter, beim nächsten Run erneut versuchen
        try:
            folder_id  = gdrive_find_or_create_folder(service, folder_name, parent_folder_id)
            folder_url = f"https://drive.google.com/drive/folders/{folder_id}"

            uploaded = 0
            for att in all_files:
                try:
                    data = gutachten_download_pdf(att["url"])
                    gdrive_upload_file(service, data, att["filename"], folder_id)
                    print(f"  [GDrive] ✅ Hochgeladen: {att['filename']}")
                    uploaded += 1
                    time.sleep(0.3)
                except Exception as up_exc:
                    print(f"  [GDrive] ⚠️  Upload fehlgeschlagen ({att['filename']}): {up_exc}")

            # Drive-Link in Notion speichern (verhindert erneuten Upload)
            notion.pages.update(
                page_id=page_id,
                properties={"Google Drive Link": {"url": folder_url}},
            )
            print(f"  [GDrive] 💾 Drive-Link gespeichert ({uploaded}/{len(all_files)} Dateien)")
            erledigt += 1

        except Exception as exc:
            print(f"  [GDrive] ❌ Drive-Fehler für '{folder_name}' (wird beim nächsten Run erneut versucht): {exc}")
        time.sleep(0.5)

    print(f"\n[GDrive] ✅ {erledigt}/{len(kandidaten)} Einträge verarbeitet")
    return erledigt


# =============================================================================
# GUTACHTEN – PDF-DOWNLOAD & PARSING
# =============================================================================

def gutachten_fetch_attachment_links(edikt_url: str) -> dict:
    """
    Öffnet die Edikt-Detailseite und gibt alle Anhang-Links zurück.
    Rückgabe: {"pdfs": [...], "images": [...]}
    """
    req = urllib.request.Request(
        edikt_url,
        headers={"User-Agent": "Mozilla/5.0 (compatible; EdikteMonitor/1.0)"}
    )
    with urllib.request.urlopen(req, timeout=30) as r:
        html = r.read().decode("utf-8", errors="replace")

    pattern = re.compile(
        r'href="(/edikte/ex/exedi3\.nsf/0/[^"]+\$file/([^"]+))"',
        re.IGNORECASE
    )
    pdfs   = []
    images = []
    for path, raw_fname in pattern.findall(html):
        fname = urllib.parse.unquote(raw_fname)
        full  = f"{BASE_URL}{path}"
        if fname.lower().endswith(".pdf"):
            pdfs.append({"url": full, "filename": fname})
        elif fname.lower().endswith((".jpg", ".jpeg", ".png")):
            images.append({"url": full, "filename": fname})
    return {"pdfs": pdfs, "images": images}


def gutachten_pick_best_pdf(pdfs: list) -> dict | None:
    """Wählt das wahrscheinlichste Gutachten-PDF aus der Liste."""
    preferred = ["gutachten", " g ", "sachverst", "sv-", "/g-", "g "]
    for pdf in pdfs:
        if any(kw in pdf["filename"].lower() for kw in preferred):
            return pdf
    for pdf in pdfs:
        if "anlagen" not in pdf["filename"].lower():
            return pdf
    return pdfs[0] if pdfs else None


def gutachten_download_pdf(url: str) -> bytes:
    """Lädt ein PDF herunter und gibt die Bytes zurück."""
    req = urllib.request.Request(
        url,
        headers={"User-Agent": "Mozilla/5.0 (compatible; EdikteMonitor/1.0)"}
    )
    with urllib.request.urlopen(req, timeout=60) as r:
        return r.read()


def _gb_extract_section(text: str, start_marker: str, end_marker: str) -> str:
    """Extrahiert Text zwischen zwei Markierungen."""
    start = text.lower().find(start_marker.lower())
    if start == -1:
        return ""
    end = text.lower().find(end_marker.lower(), start + len(start_marker))
    if end == -1:
        return text[start:]
    return text[start:end]


def _gb_parse_single_owner(lines: list, anteil_idx: int) -> dict:
    """
    Hilfsfunktion: Parst einen einzelnen Eigentümer ab einer ANTEIL:-Zeile.
    Gibt dict mit name, adresse, plz_ort, geb zurück.
    """
    adr_pattern = re.compile(
        r'GEB:\s*(\d{4}-\d{2}-\d{2})\s+ADR:\s*(.+?)\s{2,}(\d{4,5})\s*$',
        re.IGNORECASE
    )
    adr_no_geb  = re.compile(r'ADR:\s*(.+?)\s{2,}(\d{4,5})\s*$', re.IGNORECASE)
    adr_simple  = re.compile(r'ADR:\s*(.+)', re.IGNORECASE)

    owner = {"name": "", "adresse": "", "plz_ort": "", "geb": ""}

    for j in range(anteil_idx + 1, min(anteil_idx + 8, len(lines))):
        stripped = lines[j].strip()
        if not stripped:
            continue
        if re.match(r'^\d', stripped):         continue  # nächste ANTEIL-Zeile
        if re.match(r'^[a-z]\s+\d', stripped): continue  # "a 7321/2006 ..."
        if "GEB:" in stripped.upper():         continue
        if "ADR:" in stripped.upper():         continue
        if re.match(r'^\*+', stripped):        continue  # Trennlinie
        if re.match(r'^Seite\s+\d+\s+von\s+\d+', stripped, re.IGNORECASE): continue  # BUG 1: Seitenangabe

        owner["name"] = stripped

        # ADR-Zeile suchen (nächste Zeilen nach dem Namen)
        for k in range(j + 1, min(j + 4, len(lines))):
            adr_line = lines[k].strip()
            if not adr_line:
                continue
            m = adr_pattern.search(adr_line)
            if m:
                owner["geb"]     = m.group(1)
                owner["adresse"] = m.group(2).strip().rstrip(",")
                owner["plz_ort"] = m.group(3)
                break
            m2 = adr_no_geb.search(adr_line)
            if m2:
                owner["adresse"] = m2.group(1).strip().rstrip(",")
                owner["plz_ort"] = m2.group(2)
                break
            m3 = adr_simple.search(adr_line)
            if m3:
                adr_raw = m3.group(1).strip()
                plz_m   = re.search(r'\s+(\d{4,5})\s*$', adr_raw)
                if plz_m:
                    owner["plz_ort"] = plz_m.group(1)
                    owner["adresse"] = adr_raw[:plz_m.start()].strip().rstrip(",")
                else:
                    owner["adresse"] = adr_raw
                break
        break  # Name gefunden – fertig mit diesem Eigentümer

    return owner


def _gb_parse_owner(section_b: str) -> dict:
    """
    Parst ALLE Eigentümer aus Section B des Grundbuchs (Miteigentum möglich).

    Bei Miteigentum werden alle Namen mit ' | ' getrennt eingetragen.
    Adresse und PLZ/Ort kommen vom ersten Eigentümer (Haupteigentümer).

    Rückgabe:
      eigentümer_name    – alle Namen, z.B. "Hans Muster | Maria Muster"
      eigentümer_adresse – Adresse des ersten Eigentümers
      eigentümer_plz_ort – PLZ/Ort des ersten Eigentümers
      eigentümer_geb     – Geburtsdatum des ersten Eigentümers
    """
    lines   = section_b.splitlines()
    owners  = []

    for i, line in enumerate(lines):
        if "ANTEIL:" not in line.upper():
            continue
        owner = _gb_parse_single_owner(lines, i)
        if owner["name"]:
            owners.append(owner)

    if not owners:
        return {
            "eigentümer_name":    "",
            "eigentümer_adresse": "",
            "eigentümer_plz_ort": "",
            "eigentümer_geb":     "",
        }

    # BUG 1: Duplikate entfernen (z.B. GmbH die 22x in Grundbuch erscheint), Reihenfolge behalten
    seen_names: set = set()
    unique_owners = []
    for o in owners:
        if o["name"] not in seen_names:
            seen_names.add(o["name"])
            unique_owners.append(o)
    owners = unique_owners

    # Alle Namen zusammenführen ("Seite X von Y" wird durch seen_names-Filter bereits verhindert)
    alle_namen = " | ".join(o["name"] for o in owners)
    erster     = owners[0]

    return {
        "eigentümer_name":    alle_namen,
        "eigentümer_adresse": erster["adresse"],
        "eigentümer_plz_ort": erster["plz_ort"],
        "eigentümer_geb":     erster["geb"],
    }


def _gb_parse_creditors(section_c: str) -> tuple:
    """Parst Pfandrechtsgläubiger und Forderungsbeträge aus Section C."""
    gläubiger = []
    betrag    = ""
    lines = [l.strip() for l in section_c.splitlines() if l.strip()]
    fuer_pattern   = re.compile(r'^für\s+(.+)', re.IGNORECASE)
    betrag_pattern = re.compile(r'Hereinbringung von\s+(EUR\s+[\d\.,]+)', re.IGNORECASE)
    pfand_pattern  = re.compile(r'PFANDRECHT\s+Höchstbetrag\s+(EUR\s+[\d\.,]+)', re.IGNORECASE)
    seen = set()
    for line in lines:
        m = fuer_pattern.match(line)
        if m:
            name = m.group(1).strip().rstrip(".")
            if len(name) > 5 and name not in seen:
                gläubiger.append(name)
                seen.add(name)
        if not betrag:
            mb = betrag_pattern.search(line)
            if mb:
                betrag = mb.group(1).strip()
    if not betrag:
        for line in lines:
            mp = pfand_pattern.search(line)
            if mp:
                betrag = mp.group(1).strip()
                break
    return gläubiger, betrag


def gutachten_extract_info_llm(full_text: str) -> dict:
    """
    Extrahiert Eigentümer, Adresse, Gläubiger und Forderungsbetrag
    aus dem PDF-Text via OpenAI GPT-4o-mini.

    Gibt ein Result-Dict zurück (gleiche Struktur wie gutachten_extract_info).
    Bei Fehler oder fehlendem API-Key: leeres Dict.
    """
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key or not OPENAI_AVAILABLE:
        return {}

    # Nur die ersten 12.000 Zeichen senden – reicht für alle relevanten Infos
    # und hält die Token-Kosten niedrig (~0,002€ pro Dokument)
    text_snippet = full_text[:12000]

    prompt = """Du analysierst Texte aus österreichischen Gerichts-Gutachten für Zwangsversteigerungen.

Extrahiere genau diese Felder und antworte NUR mit validem JSON, ohne Erklärungen:

{
  "eigentümer_name": "Vollständiger Name der verpflichteten Partei (Immobilieneigentümer). Nur der Name, keine Adresse, kein Geburtsdatum. Mehrere Eigentümer mit ' | ' trennen.",
  "eigentümer_adresse": "Straße und Hausnummer der verpflichteten Partei (Wohnadresse für Briefversand, NICHT die Liegenschaftsadresse)",
  "eigentümer_plz_ort": "PLZ und Ort der verpflichteten Partei, z.B. '1010 Wien' oder 'D-88250 Weingarten'",
  "gläubiger": ["Liste der betreibenden Banken/Gläubiger. Nur echte Kreditgeber (Banken, Sparkassen, etc.). KEINE Anwälte, Gerichte, Sachverständige, Hausverwaltungen (WEG/EG/EGT), Aktenzeichen."],
  "forderung_betrag": "Forderungshöhe falls angegeben, z.B. 'EUR 150.000'"
}

Wichtige Regeln:
- 'Verpflichtete Partei' = Eigentümer/Schuldner → das ist eigentümer_name
- 'Betreibende Partei' = Gläubiger/Bank → das ist gläubiger
- Anwälte (RA, Rechtsanwalt, vertreten durch) sind KEINE Gläubiger
- Sachverständige, Hilfskräfte, Mitarbeiter des SV sind KEIN Eigentümer
- WEG, EG, EGT, EigG, Eigentümergemeinschaft sind KEINE Gläubiger
- Wenn ein Feld nicht gefunden wird: null
- Geburtsdaten NICHT im Namen mitgeben"""

    try:
        client = _OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user",   "content": text_snippet},
            ],
            temperature=0,          # deterministisch
            max_tokens=400,         # reicht für JSON-Antwort
            response_format={"type": "json_object"},
        )
        raw = response.choices[0].message.content.strip()
        data = json.loads(raw)
    except Exception as exc:
        print(f"    [LLM] ⚠️  OpenAI-Fehler: {exc}")
        return {}


    return {
        "eigentümer_name":    _str_val(data.get("eigentümer_name")),
        "eigentümer_adresse": _str_val(data.get("eigentümer_adresse")),
        "eigentümer_plz_ort": _str_val(data.get("eigentümer_plz_ort")),
        "eigentümer_geb":     "",
        "gläubiger":          _lst_val(data.get("gläubiger")),
        "forderung_betrag":   _str_val(data.get("forderung_betrag")),
    }


def gutachten_extract_info(pdf_bytes: bytes) -> dict:
    """
    Extrahiert Eigentümer, Adresse, Gläubiger und Forderungsbetrag aus dem PDF.
    Unterstützt Grundbuchauszug-Format (Kärnten-Stil) und professionelle
    Gutachten mit 'Verpflichtete Partei:'-Angabe (Wien-Stil).
    Gibt leeres Dict zurück wenn fitz nicht verfügbar ist.
    """
    if not FITZ_AVAILABLE:
        return {}

    doc      = fitz.open(stream=pdf_bytes, filetype="pdf")
    all_text = [p.get_text() for p in doc if p.get_text().strip()]
    full_text = "\n".join(all_text)

    result = {
        "eigentümer_name":    "",
        "eigentümer_adresse": "",
        "eigentümer_plz_ort": "",
        "eigentümer_geb":     "",
        "gläubiger":          [],
        "forderung_betrag":   "",
    }

    # ── Format 1: Grundbuchauszug Sektionen B / C ────────────────────────────
    sec_b = _gb_extract_section(full_text, "** B ***", "** C ***")
    if not sec_b:
        sec_b = _gb_extract_section(full_text, "** B **", "** C **")
    if sec_b:
        result.update(_gb_parse_owner(sec_b))

    sec_c = _gb_extract_section(full_text, "** C ***", "** HINWEIS ***")
    if not sec_c:
        sec_c = _gb_extract_section(full_text, "** C **", "HINWEIS")
    if sec_c:
        gl, bt = _gb_parse_creditors(sec_c)
        result["gläubiger"]        = gl
        result["forderung_betrag"] = bt

    # ── Format 2: Professionelles Gutachten (Verpflichtete Partei) ──────────
    # Suche im GESAMTEN Text – "Verpflichtete Partei" kann auf Seite 1, 5 oder
    # später stehen (nach Deckblatt/Inhaltsverzeichnis des Sachverständigen).
    #
    # Vorkommen:
    #   Format A (eine Zeile):  "Verpflichtete Partei: Name GmbH"
    #   Format B (nächste Zeile): "Verpflichtete Partei\n \nIng. Alfred ... GmbH"
    #
    # Adress-Extraktion: direkt aus dem Verpflichtete-Partei-Block, NICHT durch
    # spätere Namensuche – so wird die Wohnadresse des Eigentümers gefunden
    # (inkl. Deutschland D-XXXXX oder andere 5-stellige PLZ).

    # Hilfsfunktion: prüft ob eine Zeile eine Adresszeile ist
    # (Straße + Nummer) oder eine PLZ/Ort-Zeile
    def _ist_adresszeile(line: str) -> bool:
        """True wenn die Zeile wie eine Straße/Hausnummer aussieht."""
        return bool(re.search(
            r'(straße|gasse|weg|platz|allee|ring|zeile|gürtel|promenade|str\.|'
            r'strasse|gasse|graben|markt|anger|hof|aue|berg|dorf|'
            r'\d+[a-z]?\s*[/,]\s*\d|\s\d+[a-z]?$)',
            line, re.IGNORECASE))

    def _ist_plz_ort(line: str) -> tuple:
        """
        Gibt (plz, ort) zurück wenn die Zeile eine PLZ/Ort-Kombination ist.
        Unterstützt:
          - AT:  '1234 Wien'  oder  '1234'
          - DE:  'D-12345 Berlin'  oder  '12345 München'
          - Kombination in einer Zeile: 'Musterstraße 5, 1234 Wien'
        """
        # Deutsches Präfix: D-XXXXX
        m = re.search(r'\bD[-–]\s*(\d{5})\s+(.+)', line)
        if m:
            return m.group(1), f"D-{m.group(1)} {m.group(2).strip()}"
        # 5-stellige PLZ (Deutschland/Liechtenstein etc.)
        # Ortsname kann Bindestriche enthalten (z.B. Titisee-Neustadt, Baden-Baden)
        m = re.search(r'\b(\d{5})\s+([A-ZÄÖÜ][\w\-\s]+)', line)
        if m:
            plz = m.group(1)
            if not re.match(r'^(19|20)\d{3}$', plz):  # keine Jahreszahl
                ort = m.group(2).strip().rstrip('.,')   # trailing Satzzeichen weg
                return plz, f"{plz} {ort}"
        # 4-stellige PLZ (Österreich/Schweiz)
        m = re.search(r'\b(\d{4})\s+([A-ZÄÖÜ][\w\-\s]+)', line)
        if m:
            plz = m.group(1)
            if not re.match(r'^(19|20)\d{2}$', plz):
                ort = m.group(2).strip().rstrip('.,')   # trailing Satzzeichen weg
                return plz, f"{plz} {ort}"
        # Nur PLZ (4 oder 5 Stellen) ohne Ortsname
        m = re.search(r'\b(\d{4,5})\b', line)
        if m:
            plz = m.group(1)
            if not re.match(r'^(19|20)\d{2,3}$', plz):
                return plz, plz
        return "", ""

    if not result["eigentümer_name"]:
        # Alle Vorkommen von "Verpflichtete Partei" finden
        # Name + Adresse werden direkt aus diesem Block gelesen
        for vp_match in re.finditer(r'Verpflichtete\s+Partei', full_text, re.IGNORECASE):
            # Inline-Name direkt nach "Verpflichtete Partei: Name, Straße, PLZ Ort"
            # z.B. "Verpflichtete Partei: Firma XY GmbH, Kirchgasse 3, 6900 Bregenz"
            rest_of_line = full_text[vp_match.end():].split("\n")[0].strip().lstrip(":").strip()
            block = full_text[vp_match.end():vp_match.end() + 500]
            lines_vp = [l.strip().lstrip(":").strip() for l in block.split("\n")]
            lines_vp = [l for l in lines_vp if l]  # Leerzeilen raus

            name_candidate = ""
            adr_candidate  = ""
            plz_candidate  = ""

            # Sonderfall: alles in einer Zeile "Name, Straße, PLZ Ort"
            if rest_of_line and len(rest_of_line) > 3 and "," in rest_of_line:
                parts = [p.strip() for p in rest_of_line.split(",")]
                # Letzter Teil: PLZ Ort?
                plz, ort = _ist_plz_ort(parts[-1])
                if plz and len(parts) >= 2:
                    inline_name = parts[0].rstrip(".")
                    # BUG D: Hilfskraft/Mitarbeiter auch im Inline-Pfad filtern
                    # Prüfe sowohl den Namensteil als auch die gesamte Zeile
                    if re.search(
                            r'(Hilfskraft|Mitarbeiter[in]*)\s+(des|der)\s+(S[Vv]|Sachverst)',
                            rest_of_line, re.IGNORECASE):
                        pass  # nicht setzen, weiter zum nächsten vp_match
                    # BUG: Nur Punkte / Sonderzeichen ohne Buchstaben/Ziffern → überspringen
                    elif not any(c.isalnum() for c in inline_name):
                        pass
                    else:
                        name_candidate = inline_name
                        adr_candidate  = parts[-2].rstrip(".") if len(parts) >= 3 else ""
                        plz_candidate  = ort
                        result["eigentümer_name"]    = name_candidate
                        result["eigentümer_adresse"] = adr_candidate
                        result["eigentümer_plz_ort"] = plz_candidate
                        break

            for idx, line in enumerate(lines_vp):
                # Stopp: nächster Hauptabschnitt
                if re.match(r'^(wegen|gegen|Aktenzahl|Auftrag|Gericht|Betreibende|\d+\.)',
                            line, re.IGNORECASE):
                    break
                # Vertreter-Zeilen nie als Name nehmen
                if re.match(r'^(vertreten|durch:|RA\s|Rechtsanwalt)',
                            line, re.IGNORECASE):
                    break
                # Grundbuch-Anteil / Dateiname überspringen
                if re.match(r'^GA\s+\d', line, re.IGNORECASE):
                    continue
                if re.match(r'^\d+/\d+\s+(Anteil|EZ|KG)', line, re.IGNORECASE):
                    continue

                if not name_candidate:
                    # Erste brauchbare Zeile = Name
                    if len(line) > 3:
                        # BUG: Nur Punkte/Sonderzeichen ohne Buchstaben → kein Name
                        # Auch ".......... 2" (Punkte + Ziffer) → kein Name
                        if not any(c.isalpha() for c in line):
                            break
                        # BUG: Fragmente wie ") und Ma-" (PDF-Zeilenumbruch-Artefakt)
                        # Erkennbar: beginnt mit ) oder endet mit -
                        if re.match(r'^[)\]}>]', line) or line.rstrip().endswith('-'):
                            break
                        # BUG D: Hilfskraft/Mitarbeiter des SV nie als Name
                        # "- Frau Mag. Zuzana ..., Hilfskraft des Sachverständigen"
                        # "Frau Dipl.-Ing. ..., Mitarbeiterin des SV"
                        if re.search(
                                r'(Hilfskraft|Mitarbeiter[in]*)\s+(des|der)\s+(S[Vv]|Sachverst)',
                                line, re.IGNORECASE):
                            break
                        # BUG E: Kontextzeilen wie "(Sohn der verpflichteten Partei)" überspringen
                        if re.match(r'^\(', line) or re.search(
                                r'(Sohn|Tochter|Ehemann|Ehefrau|Partner)\s+(der|des)\s+verpflicht',
                                line, re.IGNORECASE):
                            break
                        # BUG C: Geburtsdatum aus Name entfernen (mit ODER ohne Komma)
                        # "Christine KLEMENT, geb.29.12.1975" → "Christine KLEMENT"
                        # "Dino Ceranic geb. 26.12.1995"      → "Dino Ceranic"
                        name_clean = re.sub(
                            r',?\s*geb\.?\s*\d{1,2}[.\-]\d{1,2}[.\-]\d{2,4}', '',
                            line, flags=re.IGNORECASE).strip().rstrip(",.")
                        # Auch "geb. DD.MM.YYYY" ohne Komma davor entfernen
                        name_clean = re.sub(
                            r'\s+geb\.?\s+\d{1,2}[.\-]\d{1,2}[.\-]\d{2,4}', '',
                            name_clean, flags=re.IGNORECASE).strip().rstrip(",.")
                        # BUG I: Name enthält komplette Adresse (Komma + PLZ/Straße)
                        # "AJ GmbH, Ragnitzstraße 91, 8047 Graz" → nur erster Teil
                        if "," in name_clean:
                            parts_n = [p.strip() for p in name_clean.split(",")]
                            plz_t, _ = _ist_plz_ort(parts_n[-1])
                            if plz_t or _ist_adresszeile(parts_n[-1]):
                                name_clean = parts_n[0].strip()
                        name_candidate = name_clean
                    continue

                # Nach dem Namen: Adresse + PLZ/Ort suchen
                # Zeile könnte Straße + PLZ/Ort in einer Zeile sein
                # z.B. "Kirchweg 3, 6900 Bregenz"
                if not adr_candidate:
                    inline_plz, inline_ort = _ist_plz_ort(line)
                    if inline_plz and _ist_adresszeile(line):
                        # Alles vor der PLZ = Straße
                        sm = re.match(r'^(.+?),?\s+(?:D[-–]\s*)?\d{4,5}\s+', line)
                        if sm:
                            adr_candidate = sm.group(1).strip().rstrip(".,")
                            plz_candidate = inline_ort
                            break
                # Zeile könnte reine Straße sein (ohne PLZ)
                # BUG F: Firmenbuchnummer nie als Adresse
                if re.match(r'^Firmenbuch', line, re.IGNORECASE):
                    break
                # BUG G: Geburtsdatum nie als Adresse ("Geb. 24. 9. 1967")
                if re.match(r'^[Gg]eb\.?\s*\d', line):
                    break
                if not adr_candidate and _ist_adresszeile(line):
                    adr_candidate = line.rstrip(".,")
                    continue

                # Zeile könnte PLZ/Ort sein
                plz, ort = _ist_plz_ort(line)
                if plz:
                    plz_candidate = ort
                    # Falls noch keine Straße: schauen ob PLZ+Ort in einer Zeile mit Straße
                    if not adr_candidate:
                        # Versuche Straße aus derselben Zeile zu lesen
                        # z.B. "Musterstraße 5, 6900 Bregenz"
                        street_m = re.match(
                            r'^(.+?),?\s+(?:D[-–]\s*)?\d{4,5}\s+', line)
                        if street_m and _ist_adresszeile(street_m.group(1)):
                            adr_candidate = street_m.group(1).strip().rstrip(".,")
                    break

                # Stopp wenn nächster Abschnitt beginnt
                if re.match(r'^(wegen|gegen|Aktenzahl|Auftrag|\d+\.)', line,
                            re.IGNORECASE):
                    break

            if name_candidate and len(name_candidate) > 3:
                result["eigentümer_name"]    = name_candidate
                result["eigentümer_adresse"] = adr_candidate
                result["eigentümer_plz_ort"] = plz_candidate
                break

    # Falls Name bekannt aber Adresse fehlt noch → nochmal im gesamten Text suchen
    # (Fallback für Fälle wo Adresse nicht direkt nach "Verpflichtete Partei" steht)
    if result["eigentümer_name"] and not result["eigentümer_adresse"]:
        name_start = re.escape(result["eigentümer_name"][:40])
        all_matches = list(re.finditer(name_start, full_text, re.IGNORECASE))
        for match_pos in reversed(all_matches):  # letztes Vorkommen zuerst
            search_block = full_text[match_pos.start():match_pos.start() + 500]
            lines_adr = [l.strip() for l in search_block.split("\n") if l.strip()]
            prev_line = ""
            for line in lines_adr[1:]:
                if re.match(r'^GA\s+\d', line, re.IGNORECASE):
                    continue
                if re.match(r'^\d+/\d+\s+(Anteil|EZ|KG)', line, re.IGNORECASE):
                    continue
                if re.match(r'^(wegen|gegen|Aktenzahl|Auftrag|\d+\.)', line,
                            re.IGNORECASE):
                    break
                # BUG F+G auch im Fallback: Firmenbuch/Geburtsdatum nie als Adresse
                if re.match(r'^Firmenbuch', line, re.IGNORECASE):
                    break
                if re.match(r'^[Gg]eb\.?\s*\d', line):
                    break
                plz, ort = _ist_plz_ort(line)
                if plz:
                    if prev_line and _ist_adresszeile(prev_line):
                        result["eigentümer_adresse"] = prev_line.rstrip(".,")
                    elif not prev_line or not _ist_adresszeile(prev_line):
                        # PLZ+Ort vielleicht in derselben Zeile wie Straße
                        street_m = re.match(
                            r'^(.+?),?\s+(?:D[-–]\s*)?\d{4,5}\s+', line)
                        if street_m and _ist_adresszeile(street_m.group(1)):
                            result["eigentümer_adresse"] = \
                                street_m.group(1).strip().rstrip(".,")
                    result["eigentümer_plz_ort"] = ort
                    break
                if _ist_adresszeile(line):
                    prev_line = line
                else:
                    prev_line = line
            if result["eigentümer_adresse"]:
                break

    # Gläubiger / Betreibende Partei – ebenfalls im gesamten Text suchen
    if not result["gläubiger"]:
        # Alle Betreibende-Partei-Blöcke sammeln (kann mehrere geben)
        gl_kandidaten: list[str] = []
        for bp_match in re.finditer(r'Betreibende\s+Partei', full_text, re.IGNORECASE):
            block = full_text[bp_match.end():bp_match.end() + 400]
            lines_block = [l.strip() for l in block.split("\n")]
            candidate = ""
            i = 0
            while i < len(lines_block):
                line_stripped = lines_block[i]
                if not line_stripped:
                    i += 1
                    continue
                # "vertreten durch:" → echter Name kommt DANACH (überspringen)
                if re.match(r'^vertreten\s+durch|^durch:', line_stripped, re.IGNORECASE):
                    # nächste nicht-leere Zeile ist der echte Gläubiger
                    for j in range(i + 1, min(i + 4, len(lines_block))):
                        next_line = lines_block[j].strip()
                        if next_line and not re.match(
                                r'^(gegen|Verpflichtete|wegen|Aktenzahl|\d+\.)',
                                next_line, re.IGNORECASE):
                            candidate = next_line
                            break
                    break
                # Nächster Abschnitt → stoppen
                if re.match(r'^(gegen\s+die|Verpflichtete|wegen|Aktenzahl)',
                            line_stripped, re.IGNORECASE):
                    break
                if line_stripped in (":", ""):
                    i += 1
                    continue
                candidate = line_stripped
                break

            if candidate and len(candidate) > 3:
                gl_kandidaten.append(candidate.rstrip(",."))

        # BUG 5+6: Gläubiger deduplicieren und EG/WEG-Hausverwaltungen filtern
        def _gl_normalize(name: str) -> str:
            """Entfernt FN-Nummern etc. für Duplikat-Vergleich."""
            return re.sub(r'\s*\(FN\s*\d+\w*\)', '', name, flags=re.IGNORECASE).strip()

        gl_seen_norm: set = set()
        gl_final: list[str] = []
        for gl in gl_kandidaten:
            # BUG A: führende ': ' entfernen (": Sparkasse Pöllau AG")
            gl = gl.lstrip(": ").strip()
            # BUG B: trailing ' |' und leere Segmente entfernen ("... AG |")
            gl = gl.rstrip(" |").strip()
            # Nach Bereinigung nochmal prüfen ob noch was übrig
            if not gl or len(gl) < 3:
                continue
            # Leere Pipe-Segmente entfernen ("| | & Gerichtsvollzieher" → weg)
            parts_gl = [p.strip() for p in gl.split("|")]
            parts_gl = [p.lstrip(": ").strip() for p in parts_gl]
            # BUG J: Gerichtsvollzieher, Rechtsanwalt o.ä. als alleinstehende Segmente filtern
            # Auch Punkteketten (".......... 2") und Personen-mit-Datum-Segmente entfernen
            def _gl_segment_ok(p: str) -> bool:
                if not p or len(p) <= 3:
                    return False
                if re.match(r'^(&\s*)?(Gerichtsvollzieher|Rechtsanwalt|RA\s|im\s+Zuge)', p, re.IGNORECASE):
                    return False
                if not any(c.isalpha() for c in p):  # nur Punkte/Ziffern/Symbole
                    return False
                # Personen-Segment mit Geburtsdatum z.B. "Elisabeth Schmid geb 1954-01-18"
                if re.search(r'\bgeb\s+\d{4}[-./]\d{2}[-./]\d{2}\b', p, re.IGNORECASE):
                    return False
                if re.search(r'\b(19|18)\d{2}[-./]\d{1,2}[-./]\d{1,2}\b', p):
                    return False
                return True
            parts_gl = [p for p in parts_gl if _gl_segment_ok(p)]
            gl = " | ".join(parts_gl).strip(" |")
            if not gl or len(gl) < 3:
                continue

            # BUG 6: "EG der EZ XXXX KG XXXXX" mit vollständiger Katastralangabe weglassen
            if re.match(r'^EG\s+der\s+EZ\s+\d+\s+KG\s+\d+', gl, re.IGNORECASE):
                continue
            # Eigentümergemeinschaft / Wohnungseigentumsgem. → kein Gläubiger
            if re.match(r'^(Eigentümergemeinschaft|Wohnungseigentums?gem\.?)', gl, re.IGNORECASE):
                continue
            # WEG / EG / EGT / EigG als Gläubiger filtern
            # "WEG EZ 2392 KG ...", "EGT Gemeinschaft ...", "EigG Kitzbühel"
            if re.match(r'^(WEG|EG[T]?|EigG)\b', gl, re.IGNORECASE):
                continue
            # Aktenzeichen als Gläubiger filtern ("Gemäß Aktenzeichen: 3 E 3374/24f")
            if re.match(r'^Gemäß\s+Aktenzeichen', gl, re.IGNORECASE):
                continue
            # Nur Punkte/Symbole ohne echte Buchstaben → kein Gläubiger
            if not any(c.isalpha() for c in gl):
                continue
            # Personen mit Geburtsdatum filtern – verschiedene Formate:
            # "Hermann Stöckl, 1920-03-29"  (ISO mit Bindestrichen)
            # "Elisabeth Schmid geb 1954-01-18"  (mit 'geb' Marker)
            # "Elisabeth Schmid geb. 25.3.1954"  (mit Punkt-Datum)
            if re.search(r'\b(19|18)\d{2}[-./]\d{1,2}[-./]\d{1,2}\b', gl):
                continue
            if re.search(r'\bgeb\.?\s*\d{1,2}[.\-]\d{1,2}[.\-]\d{2,4}', gl, re.IGNORECASE):
                continue
            if re.search(r'\bgeb\s+\d{4}[-./]\d{2}[-./]\d{2}\b', gl, re.IGNORECASE):
                continue
            # BUG H: Hotels/Gastronomiebetriebe ohne Bank-Charakter filtern
            if re.search(r'(Mountain Resort|Hotel|Gasthof|Pension|Wirtshaus|Betreiber\s+ROJ)',
                         gl, re.IGNORECASE):
                continue

            norm = _gl_normalize(gl)
            if norm not in gl_seen_norm:
                gl_seen_norm.add(norm)
                gl_final.append(gl)

        if gl_final:
            result["gläubiger"] = gl_final

    return result


def gutachten_enrich_notion_page(
    notion: Client,
    page_id: str,
    edikt_url: str,
) -> bool:
    """
    Hauptfunktion: Lädt das Gutachten-PDF von der Edikt-Seite,
    extrahiert Eigentümer/Gläubiger und schreibt sie in die Notion-Seite.

    Gibt True zurück wenn erfolgreich, False bei Fehler oder fehlendem PDF.
    Das Flag 'Gutachten analysiert?' wird immer gesetzt (True/False).
    """
    if not FITZ_AVAILABLE:
        print("    [Gutachten] ⚠️  PyMuPDF nicht verfügbar – überspringe PDF-Analyse")
        return False

    try:
        attachments = gutachten_fetch_attachment_links(edikt_url)
        pdfs = attachments["pdfs"]
    except Exception as exc:
        print(f"    [Gutachten] ⚠️  Fehler beim Laden der Edikt-Seite: {exc}")
        notion.pages.update(
            page_id=page_id,
            properties={"Gutachten analysiert?": {"checkbox": False}}
        )
        return False

    if not pdfs:
        # BUG 9: analysiert?=True setzen damit dieser Eintrag nicht endlos wiederholt wird
        print("    [Gutachten] ℹ️  Kein PDF-Anhang gefunden – markiere als abgeschlossen")
        notion.pages.update(
            page_id=page_id,
            properties={
                "Gutachten analysiert?": {"checkbox": True},
                "Notizen": {"rich_text": [{"text": {"content": "Kein PDF auf Edikt-Seite verfügbar"}}]},
            }
        )
        return False

    gutachten = gutachten_pick_best_pdf(pdfs)
    print(f"    [Gutachten] 📄 {gutachten['filename']}")

    try:
        pdf_bytes = gutachten_download_pdf(gutachten["url"])
    except Exception as exc:
        print(f"    [Gutachten] ⚠️  Download-Fehler: {exc}")
        notion.pages.update(
            page_id=page_id,
            properties={"Gutachten analysiert?": {"checkbox": False}}
        )
        return False

    # ── Text aus PDF extrahieren ─────────────────────────────────────────────
    try:
        doc       = fitz.open(stream=pdf_bytes, filetype="pdf")
        full_text = "\n".join(p.get_text() for p in doc if p.get_text().strip())
    except Exception as exc:
        print(f"    [Gutachten] ⚠️  PDF-Text-Fehler: {exc}")
        notion.pages.update(
            page_id=page_id,
            properties={"Gutachten analysiert?": {"checkbox": False}}
        )
        return False

    # ── Extraktion: LLM zuerst, Regex als Fallback ───────────────────────────
    info = {}
    used_llm = False
    if OPENAI_AVAILABLE and os.environ.get("OPENAI_API_KEY"):
        try:
            info = gutachten_extract_info_llm(full_text)
            if info.get("eigentümer_name") or info.get("gläubiger"):
                used_llm = True
                print("    [Gutachten] 🤖 LLM-Extraktion erfolgreich")
        except Exception as exc:
            print(f"    [Gutachten] ⚠️  LLM-Fehler: {exc}")
            info = {}

    if not used_llm:
        # Fallback: Regex-Parser (Grundbuchauszug-Format + VP-Block)
        try:
            info = gutachten_extract_info(pdf_bytes)
            print("    [Gutachten] 🔍 Regex-Fallback verwendet")
        except Exception as exc:
            print(f"    [Gutachten] ⚠️  Parse-Fehler: {exc}")
            notion.pages.update(
                page_id=page_id,
                properties={"Gutachten analysiert?": {"checkbox": False}}
            )
            return False

    # ── Notion-Properties aufbauen ───────────────────────────────────────────
    # has_owner wird nach Bereinigung gesetzt (weiter unten)
    properties: dict = {
        "Gutachten analysiert?": {"checkbox": True},
    }

    # ── Nachbereinigung: globale Hilfsfunktionen verwenden ──────────────────
    name_clean = _clean_name(info.get("eigentümer_name", ""))
    adr_clean  = _clean_adresse(info.get("eigentümer_adresse", ""))

    if name_clean:
        print(f"    [Gutachten] 👤 Eigentümer: {name_clean}")
        properties["Verpflichtende Partei"] = _rt(name_clean)

    if adr_clean:
        print(f"    [Gutachten] 🏠 Adresse: {adr_clean}")
        properties["Zustell Adresse"] = _rt(adr_clean)

    if info.get("eigentümer_plz_ort"):
        properties["Zustell PLZ/Ort"] = _rt(info["eigentümer_plz_ort"])

    # ── Betreibende Partei (Gläubiger / Bank) ──────────────────────────────
    if info.get("gläubiger"):
        gl_text = " | ".join(info["gläubiger"])
        print(f"    [Gutachten] 🏦 Gläubiger: {' | '.join(info['gläubiger'][:2])}")
        properties["Betreibende Partei"] = _rt(gl_text)

    # ── Notizen: Forderungsbetrag + PDF-Link ────────────────────────────────
    # HINWEIS: 'Langgutachten (Datei)' ist ein Notion-File-Upload-Feld und kann
    # keine externen URLs speichern → PDF-Link bleibt in Notizen.
    # has_owner basiert auf bereinigtem Name/Adresse
    has_owner = bool(name_clean or adr_clean)

    notiz_parts = []
    if info.get("forderung_betrag"):
        notiz_parts.append("Forderung: " + info["forderung_betrag"])
    notiz_parts.append(f"Gutachten-PDF: {gutachten['url']}")
    properties["Notizen"] = _rt("\n".join(notiz_parts))

    if not has_owner:
        # Gescanntes Dokument – trotzdem als analysiert markieren
        properties["Notizen"] = _rt(
            f"Gutachten-PDF: {gutachten['url']}\n"
            "(Kein Text lesbar – gescanntes Dokument)"
        )
        print("    [Gutachten] ⚠️  Kein Eigentümer gefunden (gescanntes Dokument?)")

    try:
        notion.pages.update(page_id=page_id, properties=properties)
        print("    [Gutachten] ✅ Notion aktualisiert")
    except Exception as exc:
        print(f"    [Gutachten] ⚠️  Notion-Update-Fehler: {exc}")
        return False

    return True


# =============================================================================
# NOTION
# =============================================================================

def notion_load_all_ids(notion: Client, db_id: str) -> dict[str, str]:
    """
    Lädt ALLE bestehenden Einträge aus der Notion-DB und gibt ein Dict
    {edikt_id -> page_id} zurück.

    Zusätzlich werden Einträge mit fortgeschrittener Workflow-Phase
    (z.B. 'Angeschrieben', 'Angebot', 'Gekauft') unter dem Sentinel-Wert
    "(geschuetzt)" gespeichert – der Scraper überspringt diese komplett,
    auch wenn die Hash-ID matcht. So werden bereits bearbeitete Immobilien
    niemals dupliziert oder überschrieben.

    Paginierung: Notion liefert max. 100 Ergebnisse pro Anfrage.
    """
    # Workflow-Phasen die NICHT überschrieben werden dürfen
    # (globale GESCHUETZT_PHASEN Konstante wird verwendet)

    print("[Notion] 📥 Lade alle bestehenden IDs aus der Datenbank …")
    known: dict[str, str] = {}  # edikt_id -> page_id  (oder "(geschuetzt)")
    has_more = True
    cursor = None
    page_count = 0
    geschuetzt_count = 0

    while has_more:
        kwargs: dict = {"filter": {"value": "page", "property": "object"}, "page_size": 100}
        if cursor:
            kwargs["start_cursor"] = cursor
        try:
            resp = notion.search(**kwargs)
        except Exception as exc:
            print(f"  [Notion] ⚠️  Fehler beim Laden der IDs: {exc}")
            break

        for page in resp.get("results", []):
            # Nur Pages aus unserer DB
            parent = page.get("parent", {})
            if parent.get("database_id", "").replace("-", "") != db_id.replace("-", ""):
                continue

            props = page.get("properties", {})

            # Workflow-Phase prüfen
            phase_sel = props.get("Workflow-Phase", {}).get("select") or {}
            phase = phase_sel.get("name", "")

            # Status-Feld prüfen:
            # 🔴 Rot              → IMMER echte page_id speichern (Entfall archiviert immer)
            #                       Rot hat Vorrang vor jeder Phase
            # 🟢 Grün / 🟡 Gelb  → komplett geschützt (kein Überschreiben, kein Auto-Archiv)
            status_sel = props.get("Status", {}).get("select") or {}
            status = status_sel.get("name", "")
            ist_rot        = (status == "🔴 Rot")
            # Rot hat Vorrang: auch wenn Phase geschützt wäre, zählt Rot
            ist_geschuetzt = (not ist_rot) and (phase in GESCHUETZT_PHASEN or status in ("🟢 Grün", "🟡 Gelb"))

            # Hash-ID auslesen
            hash_rt = props.get("Hash-ID / Vergleichs-ID", {}).get("rich_text", [])
            eid = ""
            if hash_rt:
                eid = hash_rt[0].get("plain_text", "").strip().lower()

            # Titel-Fingerprint für alle Einträge holen (wird unten gespeichert)
            title_rt_all = props.get("Liegenschaftsadresse", {}).get("title", [])
            title_all    = title_rt_all[0].get("plain_text", "").strip().lower() if title_rt_all else ""

            if eid:
                if ist_geschuetzt:
                    known[eid] = "(geschuetzt)"
                    geschuetzt_count += 1
                    # Auch Titel-Fingerprint mit page_id speichern – damit ein neues Edikt
                    # zur selben Immobilie (neue Hash-ID) erkannt und geupdated werden kann.
                    if title_all:
                        known[f"__titel__{title_all}"] = f"(geschuetzt_update:{page['id']})"
                elif ist_rot:
                    # Rot: Scraper legt keinen neuen Eintrag an (Duplikat-Schutz),
                    # aber die echte page_id bleibt gespeichert damit ein
                    # Entfall-Edikt die Seite archivieren kann.
                    known[eid] = page["id"]
                    geschuetzt_count += 1
                else:
                    known[eid] = page["id"]

            # Einträge OHNE Hash-ID aber MIT fortgeschrittener Phase:
            # Titel als Ersatz-Fingerprint speichern (verhindert Doppelanlage
            # bei manuell eingetragenen Immobilien ohne Hash-ID)
            elif ist_geschuetzt or ist_rot:
                if title_all:
                    if ist_geschuetzt:
                        known[f"__titel__{title_all}"] = f"(geschuetzt_update:{page['id']})"
                    else:
                        # Rot: echte ID damit Entfall immer greift
                        known[f"__titel__{title_all}"] = page["id"]
                    geschuetzt_count += 1

            page_count += 1

        has_more = resp.get("has_more", False)
        cursor   = resp.get("next_cursor")

    print(f"[Notion] ✅ {len(known)} Einträge geladen "
          f"({geschuetzt_count} geschützt, {page_count} Seiten geprüft)")
    return known


def notion_load_all_pages(notion: Client, db_id: str) -> list[dict]:
    """
    Lädt ALLE Pages aus der Notion-DB in einem einzigen Durchlauf.
    Gibt eine Liste aller Page-Objekte (mit Properties) zurück.

    Wird von Status-Sync, Bereinigung, Tote-URLs und Qualitäts-Check
    gemeinsam genutzt um mehrfache DB-Scans zu vermeiden.
    """
    print("[Notion] 📥 Lade alle Pages für Cleanup-Schritte …")
    pages: list[dict] = []
    has_more     = True
    start_cursor = None

    while has_more:
        kwargs: dict = {
            "filter": {"value": "page", "property": "object"},
            "page_size": 100,
        }
        if start_cursor:
            kwargs["start_cursor"] = start_cursor
        try:
            resp = notion.search(**kwargs)
        except Exception as exc:
            print(f"  [Notion] ⚠️  Fehler beim Laden der Pages: {exc}")
            break

        for page in resp.get("results", []):
            parent = page.get("parent", {})
            if parent.get("database_id", "").replace("-", "") != db_id.replace("-", ""):
                continue
            pages.append(page)

        has_more     = resp.get("has_more", False)
        start_cursor = resp.get("next_cursor")

    print(f"[Notion] ✅ {len(pages)} Pages geladen")
    return pages


def notion_create_eintrag(notion: Client, db_id: str, data: dict,
                          known_ids: dict | None = None) -> dict:
    """
    Legt einen neuen Eintrag in Notion an.
    Ruft die Detailseite ab, filtert nach Kategorie und befüllt alle Felder.
    Gibt den detail-Dict zurück (oder {} wenn Objekt gefiltert wurde).
    Rückgabe None bedeutet: Objekt wurde durch Kategorie-Filter ausgeschlossen
    oder ist ein Titel-Duplikat eines bereits geschützten Eintrags.
    """
    bundesland   = data.get("bundesland", "Unbekannt")
    link         = data.get("link", "")
    edikt_id     = data.get("edikt_id", "")
    beschreibung = data.get("beschreibung", "")
    typ          = data.get("type", "Versteigerung")

    # ── Detailseite abrufen ──────────────────────────────────────────────────
    detail: dict = {}
    if link:
        detail = fetch_detail(link)

    # ── Kategorie-Filter (auf Detailseite, zuverlässiger als Link-Text) ──────
    kategorie = detail.get("kategorie", "")
    if kategorie and is_excluded_by_kategorie(kategorie):
        print(f"  [Filter] ⛔ Kategorie ausgeschlossen: '{kategorie}' ({edikt_id[:8]}…)")
        return None  # Signalisiert: nicht importieren

    # ── Liegenschaftsadresse als Titel ───────────────────────────────────────
    adresse_voll = detail.get("adresse_voll", "")
    if not adresse_voll:
        datum_m = re.search(r"\((\d{2}\.\d{2}\.\d{4})\)", beschreibung)
        adresse_voll = f"{bundesland} – {datum_m.group(1) if datum_m else beschreibung[:60]}"

    titel    = adresse_voll
    objektart = kategorie or beschreibung[:200]

    # ── Titel-Duplikat-Check: selbe Adresse bereits als geschützt bekannt? ───
    # Fängt den Fall ab, dass dasselbe Objekt mit neuer edikt_id auftaucht
    # (z.B. neuer Versteigerungstermin) und bereits manuell bearbeitet wurde.
    if known_ids is not None:
        titel_key = f"__titel__{adresse_voll.strip().lower()}"
        val = known_ids.get(titel_key, "")
        if val.startswith("(geschuetzt_update:"):
            existing_page_id = val[len("(geschuetzt_update:"):-1]
            print(f"  [Notion] 🔄 Neues Edikt für bekannte Immobilie: {adresse_voll[:60]}")
            return ("__edikt_update__", existing_page_id, detail)
        elif val == "(geschuetzt)":
            # Altes Format ohne page_id – nur überspringen
            print(f"  [Notion] 🔒 Titel-Duplikat übersprungen (bereits geschützt): {adresse_voll[:60]}")
            return None

    # ── Kern-Properties (existieren garantiert in jeder Notion-DB) ───────────
    properties: dict = {
        "Liegenschaftsadresse": {
            "title": [{"text": {"content": titel[:200]}}]
        },
        "Hash-ID / Vergleichs-ID": {
            "rich_text": [{"text": {"content": edikt_id}}]
        },
        "Link": {"url": link},
        "Art des Edikts": {
            "select": {
                "name": typ if typ in ("Versteigerung", "Entfall des Termins") else "Versteigerung"
            }
        },
        "Bundesland":              {"select": {"name": bundesland}},
        "Neu eingelangt":          {"checkbox": True},
        "Automatisch importiert?": {"checkbox": True},
        "Workflow-Phase":          {"select": {"name": "🆕 Neu eingelangt"}},
        "Objektart": {
            "rich_text": [{"text": {"content": objektart[:200]}}]
        },
    }

    # ── Optionale Properties – werden einzeln hinzugefügt ────────────────────
    # Schlägt ein Feld fehl, wird nur dieses Feld übersprungen, nicht der ganze Eintrag.

    verkehrswert = detail.get("schaetzwert")
    if verkehrswert is not None:
        vk_str = f"{verkehrswert:,.2f} €".replace(",", "X").replace(".", ",").replace("X", ".")
        properties["Verkehrswert"] = {"rich_text": [{"text": {"content": vk_str}}]}

    termin_iso = detail.get("termin_iso")
    if termin_iso:
        properties["Versteigerungstermin"] = {"date": {"start": termin_iso}}

    # HINWEIS: "Verpflichtende Partei" wird NICHT hier befüllt –
    # der Gerichtsname (gericht) ist NICHT der Eigentümer.
    # Dieses Feld wird ausschließlich durch gutachten_enrich_notion_page
    # aus dem Gutachten-PDF extrahiert und eingetragen.

    plz_ort = detail.get("plz_ort", "")
    if plz_ort:
        # Vollständig: "1120 Wien" → "1120 Wien"
        properties[NOTION_PLZ_FIELD] = {
            "rich_text": [{"text": {"content": plz_ort.strip()[:100]}}]
        }

    flaeche = detail.get("flaeche_objekt") or detail.get("flaeche_grundstueck")
    if flaeche is not None:
        flaeche_str = f"{flaeche:,.2f} m²".replace(",", "X").replace(".", ",").replace("X", ".")
        properties["Fläche"] = {"rich_text": [{"text": {"content": flaeche_str}}]}

    # ── Seite anlegen – erst Kern, dann optionale Felder einzeln ─────────────
    # Strategie: Kern-Properties zuerst. Falls optionale Felder nicht existieren,
    # werden sie weggelassen und der Eintrag trotzdem angelegt.
    created_page = None
    try:
        created_page = notion.pages.create(parent={"database_id": db_id}, properties=properties)
        print(f"  [Notion] ✅ Erstellt: {titel[:80]}")
    except Exception as e:
        err_str = str(e)
        # Herausfinden welches Feld das Problem ist und es entfernen
        optional_fields = [NOTION_PLZ_FIELD, "Fläche", "Verkehrswert",
                           "Versteigerungstermin", "Verpflichtende Partei"]
        removed = []
        for field in optional_fields:
            if field in err_str and field in properties:
                del properties[field]
                removed.append(field)

        if removed:
            print(f"  [Notion] ⚠️  Felder nicht gefunden, übersprungen: {removed}")
            try:
                created_page = notion.pages.create(parent={"database_id": db_id}, properties=properties)
                print(f"  [Notion] ✅ Erstellt (ohne {removed}): {titel[:80]}")
            except Exception as e2:
                raise e2  # Wirklicher Fehler → nach oben weitergeben
        else:
            raise  # Kein bekanntes optionales Feld → nach oben weitergeben

    # Gibt (detail, page_id) zurück damit der Aufrufer das Gutachten anreichern kann
    new_page_id = created_page["id"] if created_page else None
    return detail, new_page_id


def notion_update_edikt_eintrag(
    notion: Client, page_id: str, item: dict, detail: dict
) -> None:
    """
    Aktualisiert einen bestehenden Notion-Eintrag wenn dasselbe Objekt mit
    einer neuen edikt_id erscheint (z.B. neuer Versteigerungstermin).
    Aktualisiert: Link, Hash-ID, Versteigerungstermin, Verkehrswert, Notizen.
    Schreibt NICHT die Phase oder den Status – diese bleiben unberührt.
    """
    new_eid  = item.get("edikt_id", "")
    new_link = item.get("link", "")

    props: dict = {}

    if new_link:
        props["Link"] = {"url": new_link}
    if new_eid:
        props["Hash-ID / Vergleichs-ID"] = {"rich_text": [{"text": {"content": new_eid}}]}

    termin_iso = detail.get("termin_iso")
    if termin_iso:
        props["Versteigerungstermin"] = {"date": {"start": termin_iso}}

    verkehrswert = detail.get("schaetzwert")
    if verkehrswert is not None:
        vk_str = f"{verkehrswert:,.2f} €".replace(",", "X").replace(".", ",").replace("X", ".")
        props["Verkehrswert"] = {"rich_text": [{"text": {"content": vk_str}}]}

    if props:
        try:
            notion.pages.update(page_id=page_id, properties=props)
            print(f"  [Notion] ✅ Edikt-Update gespeichert (neues Termin/Link/Wert)")
        except Exception as exc:
            print(f"  [Notion] ⚠️  Edikt-Update fehlgeschlagen: {exc}")


def notion_mark_entfall(notion: Client, page_id: str, item: dict) -> None:
    """
    Markiert ein bestehendes Notion-Objekt als 'Termin entfallen'.

    Verhalten je nach aktuellem Status/Phase:

    🟢 Grün / 🟡 Gelb  → Entfall nur vermerken, NICHT archivieren
                          (Immobilie ist relevant / gekauft / in Bearbeitung)

    🔴 Rot              → IMMER archivieren, egal welche Phase
                          (Rot = manuell abgelehnt/abgebrochen, auch in späteren Phasen)

    Bereits archiviert  → Nur Art des Edikts aktualisieren (bleibt im Archiv)

    Fortgeschrittene    → Nur Entfall vermerken, Phase bleibt erhalten
    Workflow-Phase      (gilt nur wenn Status NICHT Rot ist)

    Unbearbeitet        → Normal archivieren
    """
    # Phasen die NICHT auto-archiviert werden (manuell in Bearbeitung)
    # Gilt NUR wenn Status != 🔴 Rot
    SCHUTZ_PHASEN = {
        "🔎 In Prüfung",
        "✅ Relevant – Brief vorbereiten",
        "📩 Brief versendet",
        "📊 Gutachten analysiert",
    }

    # Aktuellen Zustand der Seite lesen
    try:
        page = notion.pages.retrieve(page_id=page_id)
        props = page.get("properties", {})
        phase    = (props.get("Workflow-Phase", {}).get("select") or {}).get("name", "")
        status   = (props.get("Status", {}).get("select") or {}).get("name", "")
        archiviert = props.get("Archiviert", {}).get("checkbox", False)
    except Exception as exc:
        print(f"  [Notion] ⚠️  Entfall: Seite konnte nicht gelesen werden: {exc}")
        return

    eid = item.get('edikt_id', '?')

    # Fall 1: Bereits archiviert → nur Art des Edikts anpassen, sonst nichts
    if archiviert:
        notion.pages.update(
            page_id=page_id,
            properties={
                "Art des Edikts": {"select": {"name": "Entfall des Termins"}},
            },
        )
        print(f"  [Notion] 🗄  Entfall im Archiv vermerkt: {eid}")
        return

    # Fall 2: Status Rot → IMMER archivieren (egal welche Phase)
    # Rot = manuell abgelehnt/abgebrochen; Phase bleibt erhalten damit
    # man später sehen kann in welcher Phase der Abbruch erfolgte.
    if status == "🔴 Rot":
        notion.pages.update(
            page_id=page_id,
            properties={
                "Art des Edikts": {"select": {"name": "Entfall des Termins"}},
                "Archiviert":     {"checkbox": True},
                # Workflow-Phase NICHT überschreiben → bleibt erhalten,
                # damit sichtbar ist in welcher Phase der Abbruch erfolgte
                "Neu eingelangt": {"checkbox": False},
            },
        )
        print(f"  [Notion] 🔴 Entfall archiviert (Status Rot, Phase '{phase}' bleibt erhalten): {eid}")
        return

    # Fall 3: Status Grün oder Gelb → relevant/aktiv in Bearbeitung → NUR vermerken
    if status in ("🟢 Grün", "🟡 Gelb"):
        notion.pages.update(
            page_id=page_id,
            properties={
                "Art des Edikts": {"select": {"name": "Entfall des Termins"}},
                "Neu eingelangt": {"checkbox": False},
            },
        )
        print(f"  [Notion] 🔒 Entfall vermerkt (Status {status} – kein Auto-Archiv): {eid}")
        return

    # Fall 4: Fortgeschrittene Phase ohne Status → nur vermerken
    if phase in SCHUTZ_PHASEN:
        notion.pages.update(
            page_id=page_id,
            properties={
                "Art des Edikts": {"select": {"name": "Entfall des Termins"}},
                "Neu eingelangt": {"checkbox": False},
            },
        )
        print(f"  [Notion] 🔒 Entfall vermerkt (Phase '{phase}' – kein Auto-Archiv): {eid}")
        return

    # Fall 5: Unbearbeitet (Neu eingelangt / kein Status) → normal archivieren
    notion.pages.update(
        page_id=page_id,
        properties={
            "Art des Edikts": {"select": {"name": "Entfall des Termins"}},
            "Archiviert":     {"checkbox": True},
            "Workflow-Phase": {"select": {"name": "🗄 Archiviert"}},
            "Neu eingelangt": {"checkbox": False},
        },
    )
    print(f"  [Notion] 🔴 Entfall archiviert: {eid}")


def notion_enrich_urls(notion: Client, db_id: str) -> int:
    """
    Findet Notion-Einträge OHNE Link-URL und versucht, über die Edikte-Suche
    einen passenden Eintrag zu finden.

    Strategie:
    1. Alle Pages aus der DB via search() laden.
    2. Falls die Seite eine Hash-ID hat → Link direkt konstruieren.
    3. Falls nicht → über Titel / Bundesland eine Freitextsuche machen.

    Gibt die Anzahl der erfolgreich ergänzten URLs zurück.
    """
    print("\n[URL-Anreicherung] 🔗 Suche nach Einträgen ohne URL …")

    enriched = 0

    # Alle Seiten via search() laden (notion-client v3 hat kein databases.query)
    pages_without_url: list[dict] = []
    has_more = True
    start_cursor = None

    while has_more:
        kwargs: dict = {
            "filter": {"value": "page", "property": "object"},
            "page_size": 100,
        }
        if start_cursor:
            kwargs["start_cursor"] = start_cursor

        try:
            resp = notion.search(**kwargs)
        except Exception as exc:
            print(f"  [URL-Anreicherung] ❌ Notion-Abfrage fehlgeschlagen: {exc}")
            break

        for page in resp.get("results", []):
            # Nur Pages aus unserer DB
            parent = page.get("parent", {})
            if parent.get("database_id", "").replace("-", "") != db_id.replace("-", ""):
                continue
            # Nur Pages ohne Link
            props    = page.get("properties", {})
            link_val = props.get("Link", {}).get("url")
            if not link_val:
                pages_without_url.append(page)

        has_more = resp.get("has_more", False)
        start_cursor = resp.get("next_cursor")

    print(f"  [URL-Anreicherung] 📋 {len(pages_without_url)} Einträge ohne URL gefunden")

    for page in pages_without_url:
        page_id = page["id"]
        props   = page.get("properties", {})

        # Hash-ID vorhanden? → Link direkt bauen
        hash_rt = props.get("Hash-ID / Vergleichs-ID", {}).get("rich_text", [])
        if hash_rt:
            edikt_id = hash_rt[0].get("plain_text", "").strip()
            if edikt_id and re.fullmatch(r"[0-9a-f]{32}", edikt_id):
                constructed_link = (
                    f"{BASE_URL}/edikte/ex/exedi3.nsf/alldoc/{edikt_id}!OpenDocument"
                )
                try:
                    notion.pages.update(
                        page_id=page_id,
                        properties={"Link": {"url": constructed_link}},
                    )
                    enriched += 1
                    print(f"  [URL-Anreicherung] ✅ Link gesetzt (Hash-ID): {edikt_id}")
                except Exception as exc:
                    print(f"  [URL-Anreicherung] ❌ Update fehlgeschlagen ({edikt_id}): {exc}")
                continue

        # Kein Hash-ID → Titel-Suche auf edikte.at
        title_rt = props.get("Liegenschaftsadresse", {}).get("title", [])
        titel = title_rt[0].get("plain_text", "") if title_rt else ""

        bl_prop = props.get("Bundesland", {}).get("select") or {}
        bundesland_name = bl_prop.get("name", "")
        bl_value = BUNDESLAENDER.get(bundesland_name, "")

        if not titel and not bl_value:
            print(f"  [URL-Anreicherung] ⚠️  Kein Titel/Bundesland für {page_id[:8]}…")
            continue

        # Suche für das Bundesland + Keyword aus dem Titel
        keyword = re.sub(r"(Wien|Niederösterreich|Burgenland|Oberösterreich|Salzburg|"
                         r"Steiermark|Kärnten|Tirol|Vorarlberg)", "", titel).strip()
        keyword = keyword[:40] if keyword else ""

        matches = _search_edikt_by_keyword(bl_value, keyword)
        if len(matches) == 1:
            candidate = matches[0]
            try:
                notion.pages.update(
                    page_id=page_id,
                    properties={
                        "Link": {"url": candidate["link"]},
                        "Hash-ID / Vergleichs-ID": {
                            "rich_text": [{"text": {"content": candidate["edikt_id"]}}]
                        },
                    },
                )
                enriched += 1
                print(
                    f"  [URL-Anreicherung] ✅ Link gefunden (Freitext): "
                    f"{candidate['edikt_id']}"
                )
            except Exception as exc:
                print(f"  [URL-Anreicherung] ❌ Update fehlgeschlagen: {exc}")
        elif len(matches) == 0:
            print(f"  [URL-Anreicherung] 🔍 Kein Treffer für '{titel[:50]}'")
        else:
            print(
                f"  [URL-Anreicherung] ❓ {len(matches)} Treffer (mehrdeutig) "
                f"für '{titel[:50]}' – übersprungen"
            )

    print(f"[URL-Anreicherung] ✅ {enriched} URLs ergänzt")
    return enriched


def _search_edikt_by_keyword(bl_value: str, keyword: str) -> list[dict]:
    """
    Interne Hilfsfunktion: Sucht auf edikte.at für ein Bundesland mit einem
    Freitext-Keyword und gibt die gefundenen Items zurück.
    """
    if not bl_value:
        return []

    query_parts = [f"([BL]=({bl_value}))"]
    if keyword:
        query_parts.append(keyword)

    params = urllib.parse.urlencode({
        "SearchView": "",
        "subf": "eex",
        "SearchOrder": "4",
        "SearchMax": "50",
        "retfields": f"~BL={bl_value}",
        "ftquery": keyword,
        "query": " ".join(query_parts),
    })
    url = f"{BASE_URL}/edikte/ex/exedi3.nsf/suchedi?{params}"

    try:
        req = urllib.request.Request(
            url,
            headers={"User-Agent": "Mozilla/5.0 (compatible; EdikteMonitor/1.0)"}
        )
        with urllib.request.urlopen(req, timeout=20) as r:
            html = r.read().decode("utf-8", errors="replace")
    except Exception:
        return []

    rel_pattern = re.compile(
        r'<a[^>]+href="(alldoc/([0-9a-f]+)!OpenDocument)"[^>]*>([^<]+)</a>',
        re.IGNORECASE
    )

    results = []
    for href_rel, edikt_id, link_text in rel_pattern.findall(html):
        link_text = link_text.strip()
        if not any(link_text.startswith(t) for t in RELEVANT_TYPES):
            continue
        results.append({
            "edikt_id": edikt_id.lower(),
            "link": f"{BASE_URL}/edikte/ex/exedi3.nsf/{href_rel}",
            "beschreibung": link_text,
        })
    return results


def notion_enrich_gutachten(notion: Client, db_id: str) -> int:
    """
    Findet alle Notion-Einträge die:
      - eine URL (Link) haben, UND
      - 'Gutachten analysiert?' = False / nicht gesetzt haben, UND
      - NICHT in einer geschützten Workflow-Phase sind

    Für jeden solchen Eintrag wird das Gutachten-PDF heruntergeladen
    und die Properties (Eigentümer, Adresse, Gläubiger, Forderung) befüllt.

    Das ist der Weg für manuell eingetragene Immobilien:
    Sobald die URL gesetzt wird (entweder vom Nutzer oder durch URL-Anreicherung),
    wird das Gutachten automatisch beim nächsten Lauf analysiert.

    Gibt die Anzahl der erfolgreich angereicherten Einträge zurück.
    """
    # globale GESCHUETZT_PHASEN Konstante wird verwendet

    print("\n[Gutachten-Anreicherung] 📄 Suche nach Einträgen ohne Gutachten-Analyse …")

    to_enrich: list[dict] = []
    has_more     = True
    start_cursor = None

    while has_more:
        kwargs: dict = {
            "filter": {"value": "page", "property": "object"},
            "page_size": 100,
        }
        if start_cursor:
            kwargs["start_cursor"] = start_cursor

        try:
            resp = notion.search(**kwargs)
        except Exception as exc:
            print(f"  [Gutachten-Anreicherung] ❌ Notion-Abfrage fehlgeschlagen: {exc}")
            break

        for page in resp.get("results", []):
            parent = page.get("parent", {})
            if parent.get("database_id", "").replace("-", "") != db_id.replace("-", ""):
                continue

            props = page.get("properties", {})

            # Nur Einträge in nicht-geschützter Phase
            phase = (props.get("Workflow-Phase", {}).get("select") or {}).get("name", "")
            if phase in GESCHUETZT_PHASEN:
                continue

            # Muss eine URL haben
            link_val = props.get("Link", {}).get("url")
            if not link_val:
                continue

            # Noch nicht analysiert
            analysiert = props.get("Gutachten analysiert?", {}).get("checkbox", False)
            if analysiert:
                continue

            to_enrich.append({"page_id": page["id"], "link": link_val})

        has_more     = resp.get("has_more", False)
        start_cursor = resp.get("next_cursor")

    MAX_PER_RUN = 100  # Begrenzung: max. 100 PDFs pro Run (~15–20 Min. Laufzeit)
    total_found = len(to_enrich)
    if total_found > MAX_PER_RUN:
        print(f"  [Gutachten-Anreicherung] ⚠️  {total_found} gefunden – verarbeite nur die ersten {MAX_PER_RUN} (Rest beim nächsten Run)")
        to_enrich = to_enrich[:MAX_PER_RUN]

    print(f"  [Gutachten-Anreicherung] 📋 {len(to_enrich)} Einträge werden jetzt analysiert")

    enriched = 0
    for entry in to_enrich:
        try:
            ok = gutachten_enrich_notion_page(notion, entry["page_id"], entry["link"])
            if ok:
                enriched += 1
        except Exception as exc:
            print(f"  [Gutachten-Anreicherung] ❌ Fehler für {entry['page_id'][:8]}…: {exc}")
        time.sleep(0.3)   # kurze Pause um API-Limits zu schonen

    remaining = total_found - len(to_enrich)
    if remaining > 0:
        print(f"  [Gutachten-Anreicherung] ℹ️  Noch {remaining} Einträge offen – werden in nächsten Runs verarbeitet")
    print(f"[Gutachten-Anreicherung] ✅ {enriched} Gutachten analysiert")
    return enriched


def notion_reset_falsche_verpflichtende(notion: Client, db_id: str,
                                       all_pages: list[dict] | None = None) -> int:
    """
    Einmalige Bereinigung: Findet Einträge deren 'Verpflichtende Partei'
    einen Gerichtsnamen enthält (z.B. "BG Schwaz (870)", "BG Innere Stadt Wien (001)").

    Diese Einträge wurden irrtümlich mit dem Gericht statt dem Eigentümer befüllt.

    Aktion:
      - 'Verpflichtende Partei' → leer
      - 'Gutachten analysiert?'  → False  (damit der nächste Run sie neu verarbeitet)

    Gibt die Anzahl der bereinigten Einträge zurück.
    """
    # globale GESCHUETZT_PHASEN Konstante wird verwendet

    # Gerichts-Muster: "BG Irgendwas (123)" oder "BG Irgendwas"
    GERICHT_RE = re.compile(
        r'^(BG |Bezirksgericht |LG |Landesgericht |HG |Handelsgericht )',
        re.IGNORECASE
    )

    print("\n[Bereinigung] 🔧 Suche nach Einträgen mit falschem Gericht in 'Verpflichtende Partei' …")

    pages = all_pages if all_pages is not None else notion_load_all_pages(notion, db_id)
    to_fix: list[str] = []

    for page in pages:
        props = page.get("properties", {})

        # Geschützte Phasen auslassen
        phase = (props.get("Workflow-Phase", {}).get("select") or {}).get("name", "")
        if phase in GESCHUETZT_PHASEN:
            continue

        # 'Verpflichtende Partei' lesen
        vp_rt = props.get("Verpflichtende Partei", {}).get("rich_text", [])
        vp_text = "".join(t.get("text", {}).get("content", "") for t in vp_rt).strip()

        if not vp_text:
            continue

        # Enthält der Wert einen Gerichtsnamen?
        if GERICHT_RE.match(vp_text):
            to_fix.append(page["id"])


    # Zweiter Pass: Einträge mit analysiert?=True aber OHNE Adresse → neu analysieren
    # NUR einmalig: dieser Pass wird NICHT wiederholt wenn das PDF gescannt ist.
    # Erkennungskriterium: Notizen enthält bereits "Kein PDF" oder "gescannt"
    # → diese werden NICHT zurückgesetzt (sonst Endlosschleife)
    to_reanalyze: list[str] = []
    for page in pages:  # 'pages' wurde oben bereits geladen (all_pages oder eigener Scan)
        props = page.get("properties", {})
        phase = (props.get("Workflow-Phase", {}).get("select") or {}).get("name", "")
        if phase in GESCHUETZT_PHASEN:
            continue
        # Nur Einträge die bereits als analysiert markiert sind
        analysiert = props.get("Gutachten analysiert?", {}).get("checkbox", False)
        if not analysiert:
            continue
        # Aber OHNE Zustelladresse
        adr_rt = props.get("Zustell Adresse", {}).get("rich_text", [])
        adr_text = "".join(t.get("text", {}).get("content", "") for t in adr_rt).strip()
        if not adr_text:
            # STOPP: wenn Notizen bereits "Kein PDF" oder ähnliches enthalten
            # → das PDF ist gescannt/nicht lesbar → NICHT nochmal versuchen
            notiz_rt = props.get("Notizen", {}).get("rich_text", [])
            notiz_text = "".join(t.get("text", {}).get("content", "") for t in notiz_rt).strip()
            if any(marker in notiz_text for marker in (
                "Kein PDF", "gescannt", "nicht lesbar", "kein Eigentümer"
            )):
                continue  # gescanntes Dokument → kein Reset, verhindert Endlosschleife
            # Nur zurücksetzen wenn ein Link vorhanden (sonst kein PDF zum analysieren)
            link_rt = props.get("Link", {}).get("url") or ""
            if link_rt and page["id"] not in to_fix:
                to_reanalyze.append(page["id"])

    if to_reanalyze:
        print(f"  [Bereinigung] 🔄 {len(to_reanalyze)} analysierte Einträge ohne Adresse → werden neu analysiert …")
        for page_id in to_reanalyze:
            try:
                notion.pages.update(
                    page_id=page_id,
                    properties={"Gutachten analysiert?": {"checkbox": False}}
                )
            except Exception as exc:
                print(f"  [Bereinigung] ⚠️  Fehler für {page_id[:8]}…: {exc}")
            time.sleep(0.2)

    if not to_fix and not to_reanalyze:
        print("  [Bereinigung] ✅ Keine falschen Einträge gefunden – alles in Ordnung")
        return 0

    print(f"  [Bereinigung] 🔧 {len(to_fix)} Einträge mit Gerichtsname gefunden – werden bereinigt …")

    fixed = 0
    for page_id in to_fix:
        try:
            notion.pages.update(
                page_id=page_id,
                properties={
                    "Verpflichtende Partei": {"rich_text": []},
                    "Gutachten analysiert?": {"checkbox": False},
                }
            )
            fixed += 1
        except Exception as exc:
            print(f"  [Bereinigung] ⚠️  Fehler für {page_id[:8]}…: {exc}")
        time.sleep(0.2)

    print(f"[Bereinigung] ✅ {fixed} Gerichtsname-Einträge + {len(to_reanalyze)} adresslose Einträge zurückgesetzt")
    return fixed + len(to_reanalyze)


# =============================================================================
# STATUS-SYNC – Status (Rot/Gelb/Grün) → Phase + Checkboxen automatisch setzen
# =============================================================================

def notion_status_sync(notion: Client, db_id: str,
                        all_pages: list[dict] | None = None) -> int:
    """
    Synchronisiert zwei manuelle Felder → Workflow-Phase + Checkboxen.

    ── Quelle 1: Status-Farbe ──────────────────────────────────────────────
      🔴 Rot  → Phase: '❌ Nicht relevant', Neu eingelangt: False,
                Relevanz geprüft?: True, Archiviert: True
      🟡 Gelb → Phase: '🔎 In Prüfung',   Neu eingelangt: False
      🟢 Grün → Phase: '✅ Gekauft',       Neu eingelangt: False

    ── Quelle 2: 'Für uns relevant?' (Select) ──────────────────────────────
      Ja         → Phase: '✅ Relevant – Brief vorbereiten',
                   Relevanz geprüft?: True, Neu eingelangt: False
      Nein       → Phase: '❌ Nicht relevant', Status: 🔴 Rot,
                   Relevanz geprüft?: True, Neu eingelangt: False, Archiviert: True
      Beobachten → Phase: '🔎 In Prüfung',
                   Relevanz geprüft?: True, Neu eingelangt: False

    all_pages: vorgeladene Pages (von notion_load_all_pages). Falls None,
               wird ein eigener Scan durchgeführt.
    Gibt die Anzahl aktualisierter Einträge zurück.
    """

    # Erwartete Phase je Status-Farbe
    STATUS_SOLL_PHASE = {
        "🔴 Rot":  "❌ Nicht relevant",
        "🟡 Gelb": "🔎 In Prüfung",
        "🟢 Grün": "✅ Gekauft",
    }

    # Erwartete Phase je 'Für uns relevant?'-Wert
    RELEVANT_SOLL_PHASE = {
        "Ja":         "✅ Relevant – Brief vorbereiten",
        "Nein":       "❌ Nicht relevant",
        "Beobachten": "🔎 In Prüfung",
    }

    print("\n[Status-Sync] 🔄 Prüfe Status + Relevanz → Phase …")

    pages = all_pages if all_pages is not None else notion_load_all_pages(notion, db_id)
    to_update: list[dict] = []

    for page in pages:
        props     = page.get("properties", {})
        status    = (props.get("Status", {}).get("select") or {}).get("name", "")
        relevant  = (props.get("Für uns relevant?", {}).get("select") or {}).get("name", "")
        phase_ist = (props.get("Workflow-Phase", {}).get("select") or {}).get("name", "")
        update_props: dict = {}

        # ── Quelle 2: 'Für uns relevant?' hat Vorrang vor Status-Farbe ──
        if relevant in RELEVANT_SOLL_PHASE:
            phase_soll = RELEVANT_SOLL_PHASE[relevant]

            # Immer: Relevanz geprüft + Neu eingelangt
            update_props["Relevanz geprüft?"] = {"checkbox": True}
            update_props["Neu eingelangt"]   = {"checkbox": False}

            # Phase nur setzen wenn noch nicht korrekt
            if phase_ist != phase_soll:
                update_props["Workflow-Phase"] = {"select": {"name": phase_soll}}

            # Bei Nein: zusätzlich Status Rot + Archiviert
            if relevant == "Nein":
                update_props["Status"]    = {"select": {"name": "🔴 Rot"}}
                update_props["Archiviert"] = {"checkbox": True}

        # ── Quelle 1: Status-Farbe (nur wenn kein Relevanz-Wert gesetzt) ─
        elif status in STATUS_SOLL_PHASE:
            phase_soll = STATUS_SOLL_PHASE[status]

            if phase_ist != phase_soll:
                update_props["Workflow-Phase"] = {"select": {"name": phase_soll}}

            update_props["Neu eingelangt"] = {"checkbox": False}

            if status == "🔴 Rot":
                update_props["Relevanz geprüft?"] = {"checkbox": True}
                update_props["Archiviert"]        = {"checkbox": True}

        # Keine relevanten Felder gesetzt → überspringen
        if not update_props:
            continue

        to_update.append({
            "page_id":      page["id"],
            "update_props": update_props,
            "label":        f"relevant={relevant or '–'} status={status or '–'} → phase={update_props.get('Workflow-Phase', {}).get('select', {}).get('name', phase_ist)}",
        })

    print(f"  [Status-Sync] 📋 {len(to_update)} Einträge werden synchronisiert")

    updated = 0
    for entry in to_update:
        try:
            notion.pages.update(page_id=entry["page_id"], properties=entry["update_props"])
            print(f"  [Status-Sync] ✅ {entry['label']}")
            updated += 1
        except Exception as exc:
            print(f"  [Status-Sync] ⚠️  Update fehlgeschlagen: {exc}")
        time.sleep(0.2)

    print(f"[Status-Sync] ✅ {updated} Einträge synchronisiert")
    return updated


# =============================================================================
# SCHRITT 1: QUALITÄTS-CHECK – alle analysierten Einträge auf Vollständigkeit
# =============================================================================

def notion_qualitaetscheck(notion: Client, db_id: str,
                           all_pages: list[dict] | None = None) -> int:
    """
    Geht alle Einträge durch die bereits als 'Gutachten analysiert?' = True
    markiert sind, aber eines oder mehrere dieser Felder LEER haben:
      - Verpflichtende Partei (Eigentümer)
      - Zustell Adresse
      - Betreibende Partei (Gläubiger)

    Solche Einträge werden zurückgesetzt (analysiert? = False) damit
    notion_enrich_gutachten sie beim nächsten Schritt neu analysiert.

    Einträge mit 'gescanntes Dokument' oder 'Kein PDF' im Notizen-Feld
    werden NICHT zurückgesetzt (da kein PDF vorhanden bzw. nicht lesbar).

    Gibt die Anzahl zurückgesetzter Einträge zurück.
    """
    # globale GESCHUETZT_PHASEN Konstante wird verwendet

    print("\n[Qualitäts-Check] 🔍 Prüfe alle analysierten Einträge auf Vollständigkeit …")

    pages = all_pages if all_pages is not None else notion_load_all_pages(notion, db_id)
    to_reset: list[str] = []
    total_checked = 0

    for page in pages:
        props = page.get("properties", {})

        # Nur analysierte Einträge
        analysiert = props.get("Gutachten analysiert?", {}).get("checkbox", False)
        if not analysiert:
            continue

        # Geschützte Phasen überspringen
        phase = (props.get("Workflow-Phase", {}).get("select") or {}).get("name", "")
        if phase in GESCHUETZT_PHASEN:
            continue

        # Archivierte überspringen
        archiviert = props.get("Archiviert", {}).get("checkbox", False)
        if archiviert:
            continue

        # Muss eine URL haben (sonst gibt es nichts zu analysieren)
        link_val = props.get("Link", {}).get("url")
        if not link_val:
            continue

        total_checked += 1

        # Notizen prüfen – gescannte/fehlende PDFs nicht nochmal versuchen
        notizen_rt = props.get("Notizen", {}).get("rich_text", [])
        notizen_text = "".join(
            (b.get("text") or {}).get("content", "") for b in notizen_rt
        ).lower()
        if "gescannt" in notizen_text or "kein pdf" in notizen_text or "nicht lesbar" in notizen_text:
            continue

        # Felder prüfen
        eigentümer_rt = props.get("Verpflichtende Partei", {}).get("rich_text", [])
        eigentümer    = "".join(
            (b.get("text") or {}).get("content", "") for b in eigentümer_rt
        ).strip()

        adresse_rt = props.get("Zustell Adresse", {}).get("rich_text", [])
        adresse    = "".join(
            (b.get("text") or {}).get("content", "") for b in adresse_rt
        ).strip()

        gläubiger_rt = props.get("Betreibende Partei", {}).get("rich_text", [])
        gläubiger    = "".join(
            (b.get("text") or {}).get("content", "") for b in gläubiger_rt
        ).strip()

        # Zurücksetzen wenn Eigentümer UND Adresse fehlen (beide leer)
        if not eigentümer and not adresse:
            to_reset.append(page["id"])

    print(f"  [Qualitäts-Check] 📊 {total_checked} analysierte Einträge geprüft")
    print(f"  [Qualitäts-Check] 🔄 {len(to_reset)} unvollständige Einträge → werden neu analysiert")

    reset_count = 0
    for page_id in to_reset:
        try:
            notion.pages.update(
                page_id=page_id,
                properties={"Gutachten analysiert?": {"checkbox": False}}
            )
            reset_count += 1
        except Exception as exc:
            print(f"  [Qualitäts-Check] ⚠️  Reset fehlgeschlagen für {page_id[:8]}…: {exc}")
        time.sleep(0.15)

    print(f"[Qualitäts-Check] ✅ {reset_count} Einträge zurückgesetzt")
    return reset_count


# =============================================================================
# SCHRITT 2: VISION-ANALYSE – gescannte PDFs mit GPT-4o-Vision
# =============================================================================

def gutachten_extract_info_vision(pdf_bytes: bytes, pdf_url: str) -> dict:
    """
    Analysiert ein gescanntes PDF (kein extrahierbarer Text) mit GPT-4o-Vision.
    Konvertiert die ersten 3 Seiten des PDFs in Bilder (base64) und sendet
    sie an die OpenAI Vision API.

    Gibt das gleiche Result-Dict zurück wie gutachten_extract_info_llm.
    Gibt leeres Dict zurück bei Fehler.
    """
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key or not OPENAI_AVAILABLE:
        return {}
    if not FITZ_AVAILABLE:
        return {}

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:
        print(f"    [Vision] ⚠️  PDF öffnen fehlgeschlagen: {exc}")
        return {}

    # Erste 8 Seiten als Bilder rendern – Eigentümer steht oft erst auf Seite 4–8
    # 2.5x Zoom = ~190 DPI → bessere Lesbarkeit für gescannte Dokumente
    images_b64: list[str] = []
    for page_num in range(min(8, len(doc))):
        try:
            page = doc[page_num]
            mat  = fitz.Matrix(2.5, 2.5)   # 2.5x Zoom = ~190 DPI
            pix  = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            img_bytes = pix.tobytes("jpeg", jpg_quality=80)
            images_b64.append(base64.b64encode(img_bytes).decode("utf-8"))
        except Exception as exc:
            print(f"    [Vision] ⚠️  Seite {page_num+1} konnte nicht gerendert werden: {exc}")
            continue

    if not images_b64:
        print("    [Vision] ⚠️  Keine Seiten gerendert")
        return {}

    prompt = """Du analysierst Bilder aus österreichischen Gerichts-Gutachten für Zwangsversteigerungen.
Es gibt zwei Dokumenttypen – analysiere BEIDE:

1. Professionelles Gutachten (Wien-Stil): Enthält Abschnitte 'Verpflichtete Partei' (= Eigentümer) und 'Betreibende Partei' (= Gläubiger).
2. Grundbuchauszug (Kärnten-Stil): Enthält Abschnitte '** B **' oder 'B-Blatt' (= Eigentümer mit Anteilen) und '** C **' oder 'C-Blatt' (= Pfandrechte/Gläubiger). Der Eigentümer steht nach 'Eigentumsrecht' oder 'Anteil' in Sektion B.

Extrahiere genau diese Felder und antworte NUR mit validem JSON, ohne Erklärungen:

{
  "eigentümer_name": "Vollständiger Name des Immobilieneigentümers. Nur der Name, keine Adresse, kein Geburtsdatum. Mehrere Eigentümer mit ' | ' trennen.",
  "eigentümer_adresse": "Straße und Hausnummer des Eigentümers (Wohnadresse für Briefversand, NICHT die Liegenschaftsadresse)",
  "eigentümer_plz_ort": "PLZ und Ort des Eigentümers, z.B. '1010 Wien'",
  "gläubiger": ["Liste der betreibenden Banken/Gläubiger. Nur echte Kreditgeber (Banken, Sparkassen, Raiffeisen etc.). KEINE Anwälte, Gerichte, WEG/EG/Hausverwaltungen."],
  "forderung_betrag": "Forderungshöhe falls angegeben, z.B. 'EUR 150.000'"
}

Wichtige Regeln:
- Sachverständige, Hilfskräfte des SV, Anwälte sind KEINE Eigentümer
- WEG, EG, EGT, Eigentümergemeinschaft sind KEINE Gläubiger
- Wenn ein Feld nicht gefunden wird: null"""

    # Nachricht mit allen Seiten-Bildern zusammenbauen
    content: list[dict] = [{"type": "text", "text": "Analysiere dieses Gutachten:"}]
    for img_b64 in images_b64:
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/jpeg;base64,{img_b64}",
                "detail": "high"
            }
        })

    try:
        client   = _OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o",            # Vision-fähiges Modell (nicht mini!)
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user",   "content": content},
            ],
            temperature=0,
            max_tokens=500,
            response_format={"type": "json_object"},
        )
        raw  = response.choices[0].message.content.strip()
        data = json.loads(raw)
        print(f"    [Vision] 🔭 GPT-4o Vision analysiert ({len(images_b64)} Seiten)")
    except Exception as exc:
        print(f"    [Vision] ⚠️  OpenAI Vision-Fehler: {exc}")
        return {}

    return {
        "eigentümer_name":    _str_val(data.get("eigentümer_name")),
        "eigentümer_adresse": _str_val(data.get("eigentümer_adresse")),
        "eigentümer_plz_ort": _str_val(data.get("eigentümer_plz_ort")),
        "eigentümer_geb":     "",
        "gläubiger":          _lst_val(data.get("gläubiger")),
        "forderung_betrag":   _str_val(data.get("forderung_betrag")),
    }


def notion_enrich_gescannte(notion: Client, db_id: str) -> int:
    """
    Findet alle Einträge die als 'gescanntes Dokument' markiert sind
    (Notizen enthält 'gescanntes Dokument' oder 'Kein Text lesbar')
    und versucht sie mit GPT-4o Vision neu zu analysieren.

    Gibt die Anzahl erfolgreich analysierter Einträge zurück.
    """
    if not OPENAI_AVAILABLE or not os.environ.get("OPENAI_API_KEY"):
        print("[Vision-Analyse] ℹ️  Kein OpenAI API-Key – überspringe Vision-Analyse")
        return 0
    if not FITZ_AVAILABLE:
        print("[Vision-Analyse] ℹ️  PyMuPDF nicht verfügbar – überspringe Vision-Analyse")
        return 0

    # globale GESCHUETZT_PHASEN Konstante wird verwendet

    print("\n[Vision-Analyse] 🔭 Suche nach gescannten PDFs …")

    to_vision: list[dict] = []
    has_more     = True
    start_cursor = None

    while has_more:
        kwargs: dict = {
            "filter": {"value": "page", "property": "object"},
            "page_size": 100,
        }
        if start_cursor:
            kwargs["start_cursor"] = start_cursor

        try:
            resp = notion.search(**kwargs)
        except Exception as exc:
            print(f"  [Vision-Analyse] ❌ Notion-Abfrage fehlgeschlagen: {exc}")
            break

        for page in resp.get("results", []):
            parent = page.get("parent", {})
            if parent.get("database_id", "").replace("-", "") != db_id.replace("-", ""):
                continue

            props = page.get("properties", {})

            # Nur analysierte Einträge
            analysiert = props.get("Gutachten analysiert?", {}).get("checkbox", False)
            if not analysiert:
                continue

            # Geschützte Phasen + Archivierte überspringen
            phase = (props.get("Workflow-Phase", {}).get("select") or {}).get("name", "")
            if phase in GESCHUETZT_PHASEN:
                continue
            if props.get("Archiviert", {}).get("checkbox", False):
                continue

            # Muss URL haben
            link_val = props.get("Link", {}).get("url")
            if not link_val:
                continue

            # Notizen prüfen: enthält 'gescanntes Dokument' oder 'Kein Text lesbar'?
            notizen_rt = props.get("Notizen", {}).get("rich_text", [])
            notizen_text = "".join(
                (b.get("text") or {}).get("content", "") for b in notizen_rt
            )
            # Marker für gescannte Dokumente (original oder nach Vision-Versuch)
            ist_gescannt = (
                "gescannt" in notizen_text.lower()
                or "kein text lesbar" in notizen_text.lower()
                or "via gpt-4o vision" in notizen_text.lower()
                or "unleserlich" in notizen_text.lower()
            )
            if not ist_gescannt:
                continue

            # PDF-URL aus Notizen extrahieren
            pdf_url_match = re.search(r'Gutachten-PDF:\s*(https?://\S+)', notizen_text)
            pdf_url = pdf_url_match.group(1).strip() if pdf_url_match else None

            # Eigentümer noch leer?
            eigentümer_rt = props.get("Verpflichtende Partei", {}).get("rich_text", [])
            eigentümer    = "".join(
                (b.get("text") or {}).get("content", "") for b in eigentümer_rt
            ).strip()
            if eigentümer:
                continue  # Eigentümer bereits vorhanden – überspringen

            to_vision.append({
                "page_id": page["id"],
                "link":    link_val,
                "pdf_url": pdf_url,
                "notizen": notizen_text,
            })

        has_more     = resp.get("has_more", False)
        start_cursor = resp.get("next_cursor")

    MAX_VISION = 20   # GPT-4o ist teurer → max 20 pro Run (~0.40€)
    total_found = len(to_vision)
    if total_found > MAX_VISION:
        print(f"  [Vision-Analyse] ⚠️  {total_found} gefunden – verarbeite nur die ersten {MAX_VISION}")
        to_vision = to_vision[:MAX_VISION]

    print(f"  [Vision-Analyse] 📋 {len(to_vision)} gescannte PDFs werden analysiert")

    enriched = 0
    for entry in to_vision:
        try:
            # PDF direkt laden (URL aus Notizen oder neu von Edikt-Seite holen)
            pdf_url = entry["pdf_url"]
            if not pdf_url:
                # PDF-URL neu von der Edikt-Seite laden
                try:
                    attachments = gutachten_fetch_attachment_links(entry["link"])
                    pdfs = attachments.get("pdfs", [])
                    if pdfs:
                        best = gutachten_pick_best_pdf(pdfs)
                        pdf_url = best["url"] if best else None
                except Exception as exc:
                    print(f"    [Vision] ⚠️  Edikt-Seite nicht ladbar: {exc}")
                    continue

            if not pdf_url:
                print(f"    [Vision] ⚠️  Keine PDF-URL gefunden für {entry['page_id'][:8]}…")
                continue

            pdf_bytes = gutachten_download_pdf(pdf_url)
            info = gutachten_extract_info_vision(pdf_bytes, pdf_url)

            if not info.get("eigentümer_name") and not info.get("eigentümer_adresse"):
                # Als endgültig unleserlich markieren → nie wieder versuchen
                try:
                    notizen_alt = entry["notizen"].strip()
                    # Alten gescannt-Vermerk durch finalen ersetzen
                    notizen_neu = re.sub(
                        r'\(Kein Text lesbar[^)]*\)|\(Via GPT-4o Vision[^)]*\)',
                        '', notizen_alt
                    ).strip()
                    notizen_neu += "\n(Endgültig unleserlich – kein Eigentümer auffindbar)"
                    notion.pages.update(
                        page_id=entry["page_id"],
                        properties={
                            "Notizen": {"rich_text": [{"text": {"content": notizen_neu[:2000]}}]}
                        }
                    )
                except Exception:
                    pass
                print(f"    [Vision] ℹ️  Kein Eigentümer gefunden → als unleserlich markiert")
                continue

            # Notion-Properties aufbauen (globale Hilfsfunktionen)
            name_clean = _clean_name(info.get("eigentümer_name", ""))
            adr_clean  = _clean_adresse(info.get("eigentümer_adresse", ""))

            properties: dict = {"Gutachten analysiert?": {"checkbox": True}}

            if name_clean:
                print(f"    [Vision] 👤 Eigentümer: {name_clean}")
                properties["Verpflichtende Partei"] = _rt(name_clean)

            if adr_clean:
                print(f"    [Vision] 🏠 Adresse: {adr_clean}")
                properties["Zustell Adresse"] = _rt(adr_clean)

            if info.get("eigentümer_plz_ort"):
                properties["Zustell PLZ/Ort"] = _rt(info["eigentümer_plz_ort"])

            if info.get("gläubiger"):
                gl_text = " | ".join(info["gläubiger"])
                print(f"    [Vision] 🏦 Gläubiger: {gl_text[:80]}")
                properties["Betreibende Partei"] = _rt(gl_text)

            # Notizen aktualisieren (gescannt-Vermerk entfernen)
            notiz_parts = []
            if info.get("forderung_betrag"):
                notiz_parts.append("Forderung: " + info["forderung_betrag"])
            notiz_parts.append(f"Gutachten-PDF: {pdf_url}")
            notiz_parts.append("(Via GPT-4o Vision analysiert – gescanntes Dokument)")
            properties["Notizen"] = _rt("\n".join(notiz_parts))

            notion.pages.update(page_id=entry["page_id"], properties=properties)
            print(f"    [Vision] ✅ Notion aktualisiert")
            enriched += 1

        except Exception as exc:
            print(f"  [Vision-Analyse] ❌ Fehler für {entry['page_id'][:8]}…: {exc}")
        time.sleep(0.5)  # etwas mehr Pause wegen größerer API-Anfragen

    print(f"[Vision-Analyse] ✅ {enriched} gescannte PDFs erfolgreich analysiert")
    return enriched


# =============================================================================
# SCHRITT 3: TOTE URLs – HTTP 404 → automatisch archivieren
# =============================================================================

def notion_archiviere_tote_urls(notion: Client, db_id: str,
                                all_pages: list[dict] | None = None) -> tuple[int, list[str]]:
    """
    Prüft ALLE Einträge (außer bereits archivierte) auf HTTP 404.

    Archivierungs-Logik basierend auf Status und Phase:

    ┌─────────────────────────────────────┬──────────────────────────────────────┐
    │ Status = 🟢 Grün oder 🟡 Gelb       │ Nur Telegram-Alarm, KEIN Archivieren │
    │ (egal welche Phase)                 │ (aktive Bearbeitung läuft noch)       │
    ├─────────────────────────────────────┼──────────────────────────────────────┤
    │ Status leer / grau                  │ → 🗄 Archiviert                       │
    │ Phase = 📩 Brief versendet          │ → 🗄 Archiviert + Telegram-Alarm      │
    │ Alle anderen                        │ → 🗄 Archiviert (still)               │
    └─────────────────────────────────────┴──────────────────────────────────────┘

    Gibt (Anzahl archivierter Einträge, Liste der Telegram-Alarm-Texte) zurück.
    """
    # Nur wirklich fertig archivierte überspringen
    SKIP_PHASEN = {"🗄 Archiviert"}

    # Schutz-Status: bei diesen wird NUR alarmiert, nicht archiviert
    SCHUTZ_STATUS = {"🟢 Grün", "🟡 Gelb"}

    print("\n[Tote-URLs] 🔗 Prüfe URLs auf 404 …")

    pages = all_pages if all_pages is not None else notion_load_all_pages(notion, db_id)
    to_check: list[dict] = []

    for page in pages:
        props = page.get("properties", {})

        # Bereits archivierte überspringen
        phase = (props.get("Workflow-Phase", {}).get("select") or {}).get("name", "")
        if phase in SKIP_PHASEN:
            continue
        if props.get("Archiviert", {}).get("checkbox", False):
            continue

        # Muss eine URL haben
        link_val = props.get("Link", {}).get("url")
        if not link_val:
            continue

        status_val = (props.get("Status", {}).get("select") or {}).get("name", "")

        # Titel für Alarm
        titel_rt = props.get("Liegenschaftsadresse", {}).get("title", [])
        titel = "".join(
            (b.get("text") or {}).get("content", "") for b in titel_rt
        ).strip() or page["id"][:8]

        to_check.append({
            "page_id":  page["id"],
            "link":     link_val,
            "phase":    phase,
            "status":   status_val,
            "titel":    titel,
        })

    MAX_CHECK = 50   # max 50 URL-Checks pro Run (schont das Netz)
    if len(to_check) > MAX_CHECK:
        to_check = to_check[:MAX_CHECK]

    print(f"  [Tote-URLs] 📋 {len(to_check)} Einträge werden geprüft")

    archived      = 0
    alarm_lines: list[str] = []   # Telegram-Alarme für geschützte Einträge

    for entry in to_check:
        is_404 = False
        try:
            req = urllib.request.Request(
                entry["link"],
                headers={"User-Agent": "Mozilla/5.0 (compatible; EdikteMonitor/1.0)"}
            )
            with urllib.request.urlopen(req, timeout=10) as r:
                _ = r.read(1)   # nur Header laden
        except urllib.error.HTTPError as e:
            if e.code == 404:
                is_404 = True
        except Exception:
            pass  # Netzwerkfehler / Timeout → kein 404

        if not is_404:
            time.sleep(0.2)
            continue

        print(f"  [Tote-URLs] 🗑  HTTP 404: {entry['titel'][:60]} (Phase: {entry['phase']}, Status: {entry['status'] or '–'})")

        # ── Schutz-Status: nur alarmieren, NICHT archivieren ──────────────
        if entry["status"] in SCHUTZ_STATUS:
            # Notiz lesen um zu prüfen ob bereits alarmiert wurde (einmaliger Alarm)
            bereits_alarmiert = False
            notizen_alt = ""
            try:
                page_data  = notion.pages.retrieve(page_id=entry["page_id"])
                notizen_rt = page_data["properties"].get("Notizen", {}).get("rich_text", [])
                notizen_alt = "".join(
                    (b.get("text") or {}).get("content", "") for b in notizen_rt
                ).strip()
                bereits_alarmiert = "Edikt-Seite nicht mehr verfügbar" in notizen_alt
            except Exception as exc2:
                print(f"  [Tote-URLs] ⚠️  Notiz-Lesen fehlgeschlagen: {exc2}")

            if not bereits_alarmiert:
                # Erster Alarm: Telegram + Notion-Notiz setzen
                alarm_lines.append(
                    f"⚠️ Edikt verschwunden (Status {entry['status']}): "
                    f"<b>{entry['titel'][:80]}</b>\n"
                    f"<i>Bitte in Notion als gelöscht markieren.</i>"
                )
                notizen_neu = (notizen_alt + "\n" if notizen_alt else "") + \
                              "⚠️ Edikt-Seite nicht mehr verfügbar (HTTP 404) – bitte manuell prüfen"
                try:
                    notion.pages.update(
                        page_id=entry["page_id"],
                        properties={
                            "Notizen": {"rich_text": [{"text": {"content": notizen_neu[:2000]}}]},
                        }
                    )
                except Exception as exc2:
                    print(f"  [Tote-URLs] ⚠️  Notiz-Update fehlgeschlagen: {exc2}")
            else:
                print(f"  [Tote-URLs] ℹ️  Bereits alarmiert, kein erneuter Telegram-Alarm: {entry['titel'][:50]}")
            time.sleep(0.2)
            continue

        # ── Alle anderen: archivieren ──────────────────────────────────────
        # Bei "Brief versendet" zusätzlich Telegram-Alarm
        if entry["phase"] == "📩 Brief versendet":
            alarm_lines.append(
                f"📬 Brief bereits versendet – Edikt jetzt weg: "
                f"<b>{entry['titel'][:80]}</b> → archiviert"
            )

        try:
            page_data   = notion.pages.retrieve(page_id=entry["page_id"])
            notizen_rt  = page_data["properties"].get("Notizen", {}).get("rich_text", [])
            notizen_alt = "".join(
                (b.get("text") or {}).get("content", "") for b in notizen_rt
            ).strip()
            notizen_neu = (notizen_alt + "\n" if notizen_alt else "") + \
                          "Edikt-Seite nicht mehr verfügbar (HTTP 404) – automatisch archiviert"

            notion.pages.update(
                page_id=entry["page_id"],
                properties={
                    "Archiviert":    {"checkbox": True},
                    "Workflow-Phase": {"select": {"name": "🗄 Archiviert"}},
                    "Notizen":       {"rich_text": [{"text": {"content": notizen_neu[:2000]}}]},
                }
            )
            archived += 1
        except Exception as exc2:
            print(f"  [Tote-URLs] ⚠️  Archivierung fehlgeschlagen: {exc2}")

        time.sleep(0.2)

    print(f"[Tote-URLs] ✅ {archived} tote URLs archiviert")
    return archived, alarm_lines


# =============================================================================
# BRIEF-WORKFLOW – Brief erstellen für relevante Einträge
# =============================================================================
#
# Ablauf:
#   1. Suche alle Einträge mit Phase "✅ Relevant – Brief vorbereiten"
#      bei denen "Brief erstellt am" noch LEER ist.
#   2. Bestimme zuständige Person anhand des Bundeslandes.
#   3. Befülle DOCX-Vorlage (brief_vorlage.docx) mit Platzhaltern.
#   4. Konvertiere DOCX → PDF (via reportlab/python-docx).
#   5. Lade PDF als GitHub-Artifact hoch ODER schreibe Pfad in Notizen.
#   6. Setze "Brief erstellt am" in Notion (heutiges Datum).
#   7. Sende Telegram-Nachricht mit Zusammenfassung.
#
# Kontaktdaten der Zuständigen:
#   ┌─────────────────────────────────────────────────────────────────┐
#   │ Bundesland        │ Name         │ Tel          │ E-Mail        │
#   ├───────────────────┼──────────────┼──────────────┼───────────────┤
#   │ Wien, Steiermark  │ Benjamin     │ PLACEHOLDER  │ PLACEHOLDER   │
#   │ NÖ, Burgenland    │ Christopher  │ PLACEHOLDER  │ PLACEHOLDER   │
#   │ Kärnten, Sbg, OÖ │ Du (Alex)    │ PLACEHOLDER  │ PLACEHOLDER   │
#   │ Tirol, Vorarlberg │ (noch offen) │ –            │ –             │
#   └─────────────────────────────────────────────────────────────────┘
#
# WICHTIG: Kontaktdaten unten in KONTAKT_DATEN eintragen!
# =============================================================================

# ── Kontaktdaten der Betreuer (Bundesland → Ansprechpartner) ─────────────────
#
# Benjamin Pippan    → Wien, Oberösterreich
# Christopher Dovjak → Niederösterreich, Burgenland
# Friedrich Prause   → Steiermark, Kärnten, Salzburg, Tirol, Vorarlberg
#
KONTAKT_DATEN: dict[str, dict] = {
    "Wien":             {"name": "Benjamin Pippan",    "tel": "+43699 133 90 251", "email": "office@benana.at"},
    "Oberösterreich":   {"name": "Benjamin Pippan",    "tel": "+43699 133 90 251", "email": "office@benana.at"},
    "Niederösterreich": {"name": "Christopher Dovjak", "tel": "+43 664 4531399",   "email": "christopher.dovjak@dp-im.at"},
    "Burgenland":       {"name": "Christopher Dovjak", "tel": "+43 664 4531399",   "email": "christopher.dovjak@dp-im.at"},
    "Steiermark":       {"name": "Friedrich Prause",   "tel": "+43 664 1843888",   "email": "friedrich.prause@dp-im.at"},
    "Kärnten":          {"name": "Friedrich Prause",   "tel": "+43 664 1843888",   "email": "friedrich.prause@dp-im.at"},
    "Salzburg":         {"name": "Friedrich Prause",   "tel": "+43 664 1843888",   "email": "friedrich.prause@dp-im.at"},
    "Tirol":            {"name": "Friedrich Prause",   "tel": "+43 664 1843888",   "email": "friedrich.prause@dp-im.at"},
    "Vorarlberg":       {"name": "Friedrich Prause",   "tel": "+43 664 1843888",   "email": "friedrich.prause@dp-im.at"},
}

BRIEF_VORLAGE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "brief_vorlage.docx")

# Prüfe ob python-docx verfügbar ist
try:
    from docx import Document as _DocxDocument
    from docx.shared import Pt as _DocxPt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _brief_fill_template(vorlage_path: str, platzhalter: dict[str, str]) -> bytes:
    """
    Lädt die DOCX-Vorlage, ersetzt alle {{PLATZHALTER}} und gibt den DOCX-
    Inhalt als Bytes zurück.

    Unterstützt sowohl normale Runs als auch Hyperlink-Paragraphen
    (bei denen der Text in w:hyperlink/w:r/w:t steckt und .runs leer ist).
    """
    from docx import Document
    from io import BytesIO

    doc = Document(vorlage_path)
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def replace_in_paragraph(para):
        """Ersetzt Platzhalter in einem Paragraphen (Runs + Hyperlinks)."""
        # --- Variante 1: normale Runs ---
        if para.runs:
            full_text = "".join(r.text for r in para.runs)
            new_text = full_text
            for key, val in platzhalter.items():
                new_text = new_text.replace(f"{{{{{key}}}}}", val)
            if new_text != full_text:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            return

        # --- Variante 2: Hyperlink-Struktur (keine Runs) ---
        t_elements = para._element.findall(f".//{{{W_NS}}}t")
        if not t_elements:
            return
        # Gesamttext aus allen w:t zusammensetzen
        full_text = "".join((t.text or "") for t in t_elements)
        new_text = full_text
        for key, val in platzhalter.items():
            new_text = new_text.replace(f"{{{{{key}}}}}", val)
        if new_text != full_text:
            # Ersten w:t mit neuem Text füllen, Rest leeren
            t_elements[0].text = new_text
            for t in t_elements[1:]:
                t.text = ""

    for para in doc.paragraphs:
        replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _geschlecht_via_gpt(vorname: str) -> str | None:
    """
    Fragt ChatGPT nach dem Geschlecht eines Vornamens.
    Gibt "m" (männlich), "f" (weiblich) oder None (unbekannt/neutral) zurück.
    Wird gecacht um API-Kosten zu minimieren.
    """
    try:
        api_key = os.getenv("OPENAI_API_KEY", "")
        if not api_key or not OPENAI_AVAILABLE:
            return None

        client = _OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{
                "role": "user",
                "content": (
                    f"Ist der Vorname \"{vorname}\" männlich oder weiblich? "
                    "Antworte NUR mit einem einzigen Buchstaben: "
                    "m (männlich), f (weiblich) oder n (neutral/unbekannt). "
                    "Keine Erklärung, nur den Buchstaben."
                )
            }],
            max_tokens=1,
            temperature=0,
        )
        antwort = response.choices[0].message.content.strip().lower()
        if antwort == "m":
            return "m"
        elif antwort == "f":
            return "f"
        return None
    except Exception as exc:
        print(f"  [Anrede] ⚠️  GPT-Geschlechtserkennung fehlgeschlagen: {exc}")
        return None


# Cache damit jeder Vorname nur einmal abgefragt wird
_geschlecht_cache: dict[str, str | None] = {}


def _brief_anrede(eigentuemer: str) -> str:
    """
    Erzeugt eine korrekte Anrede aus dem Eigentümernamen (verpflichtende Partei).

    Reihenfolge:
      1. Firma / GmbH / AG etc.           → "Sehr geehrte Damen und Herren,"
      2. Mehrere Personen (" und ", " & ") → "Sehr geehrte Damen und Herren,"
      3. "Herr" / "Hr." explizit          → "Sehr geehrter Herr [Nachname],"
      4. "Frau" / "Fr." explizit          → "Sehr geehrte Frau [Nachname],"
      5. Vorname via ChatGPT erkennen     → männlich/weiblich/neutral
      6. Fallback                         → "Sehr geehrte Damen und Herren,"
    """
    name = eigentuemer.strip()
    lower = name.lower()

    # ── 1. Firma / Mehrere Personen → neutral ────────────────────────────────
    FIRMA_KEYWORDS = (
        "gmbh", "ag ", " ag", "og ", " og", "kg ", " kg", "keg", "stiftung",
        "verein", "genossenschaft", "gbr", "inc.", "ltd", "s.r.o",
        "immobilien", "holding", "gmbh & co", "eigentümergemeinschaft",
    )
    if any(kw in lower for kw in FIRMA_KEYWORDS):
        return "Sehr geehrte Damen und Herren,"

    if any(sep in name for sep in (" und ", " & ", " / ", " u. ", " | ")):
        return "Sehr geehrte Damen und Herren,"

    # ── Titel extrahieren ────────────────────────────────────────────────────
    TITEL = ("Dr.", "Mag.", "Ing.", "DI", "Dipl.-Ing.", "Prof.", "DDr.",
             "MBA", "MSc", "BSc", "MMag.", "MAS", "LL.M.", "BEd", "MEd")
    titel_teile = []
    rest = name
    for t in TITEL:
        if rest.startswith(t + " ") or f" {t} " in rest:
            titel_teile.append(t)
            rest = rest.replace(t, "").strip()
    titel_str = " ".join(titel_teile) + " " if titel_teile else ""

    # ── 3. Explizites "Herr" / "Hr." ─────────────────────────────────────────
    if re.search(r"\bherr\b|\bhr\.", lower):
        clean    = re.sub(r"\b(herr|hr\.)\b", "", rest, flags=re.IGNORECASE).strip()
        nachname = clean.split()[-1] if clean.split() else clean
        return f"Sehr geehrter Herr {titel_str}{nachname},"

    # ── 4. Explizites "Frau" / "Fr." ─────────────────────────────────────────
    if re.search(r"\bfrau\b|\bfr\.", lower):
        clean    = re.sub(r"\b(frau|fr\.)\b", "", rest, flags=re.IGNORECASE).strip()
        nachname = clean.split()[-1] if clean.split() else clean
        return f"Sehr geehrte Frau {titel_str}{nachname},"

    # ── 5. Vorname via ChatGPT erkennen ──────────────────────────────────────
    # Annahme: Format "Nachname Vorname" oder "Vorname Nachname"
    # Wir probieren beide Varianten: erstes und letztes Wort als Vorname
    woerter = rest.split()
    nachname = woerter[-1] if woerter else rest

    # Ersten und letzten Wort als möglichen Vornamen testen
    vorname_kandidaten = []
    if len(woerter) >= 2:
        vorname_kandidaten = [woerter[0], woerter[-2]]  # erstes Wort, vorletztes Wort
    elif len(woerter) == 1:
        vorname_kandidaten = [woerter[0]]

    geschlecht = None
    vorname_gefunden = None
    for vn in vorname_kandidaten:
        if vn in _geschlecht_cache:
            geschlecht = _geschlecht_cache[vn]
        else:
            geschlecht = _geschlecht_via_gpt(vn)
            _geschlecht_cache[vn] = geschlecht
        if geschlecht in ("m", "f"):
            vorname_gefunden = vn
            break

    if geschlecht == "m":
        print(f"  [Anrede] 👤 {vorname_gefunden} → männlich")
        return f"Sehr geehrter Herr {titel_str}{nachname},"
    elif geschlecht == "f":
        print(f"  [Anrede] 👤 {vorname_gefunden} → weiblich")
        return f"Sehr geehrte Frau {titel_str}{nachname},"

    # ── 6. Fallback → neutral ────────────────────────────────────────────────
    return "Sehr geehrte Damen und Herren,"


def _brief_send_email(kontakt_email: str, kontakt_name: str,
                      eigentuemer: str, titel: str,
                      docx_bytes: bytes, dateiname_docx: str) -> bool:
    """
    Sendet den Brief als DOCX-Anhang per E-Mail via SendGrid API.

    Benötigte Umgebungsvariablen:
      SENDGRID_API_KEY  – API-Key (beginnt mit SG.)
      SMTP_USER         – Absender-Adresse (muss in SendGrid verifiziert sein)

    Gibt True bei Erfolg, False bei Fehler zurück.
    """
    api_key   = os.environ.get("SENDGRID_API_KEY", "")
    absender  = os.environ.get("SMTP_USER", "")

    if not api_key or not absender:
        print("  [Brief] ℹ️  SendGrid nicht konfiguriert (SENDGRID_API_KEY/SMTP_USER fehlt)")
        return False

    try:
        # DOCX als Base64
        docx_b64 = base64.b64encode(docx_bytes).decode("utf-8")

        body_text = "\n".join([
            f"Hallo {kontakt_name},",
            "",
            "anbei der Anschreiben-Entwurf für:",
            f"  Eigentümer:   {eigentuemer}",
            f"  Liegenschaft: {titel}",
            "",
            "Bitte ausdrucken und versenden.",
            "",
            "Automatisch erstellt vom Edikte-Monitor.",
        ])

        payload = {
            "personalizations": [{"to": [{"email": kontakt_email, "name": kontakt_name}]}],
            "from": {"email": absender},
            "subject": f"Neuer Anschreiben-Entwurf: {eigentuemer[:60]}",
            "content": [{"type": "text/plain", "value": body_text}],
            "attachments": [{
                "content":     docx_b64,
                "type":        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "filename":    dateiname_docx,
                "disposition": "attachment",
            }],
        }

        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        req  = urllib.request.Request(
            "https://api.sendgrid.com/v3/mail/send",
            data=data,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type":  "application/json",
            },
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=30) as r:
            status = r.getcode()

        if status in (200, 202):
            print(f"  [Brief] ✉️  E-Mail gesendet an {kontakt_email}")
            return True
        else:
            print(f"  [Brief] ⚠️  SendGrid HTTP {status}")
            return False

    except Exception as exc:
        print(f"  [Brief] ⚠️  E-Mail-Versand fehlgeschlagen: {exc}")
        return False


def notion_brief_erstellen(notion: "Client", db_id: str,
                            all_pages: list[dict] | None = None) -> tuple[int, list[str]]:
    """
    Erstellt Briefe für alle Einträge mit Phase '✅ Relevant – Brief vorbereiten'
    bei denen 'Brief erstellt am' noch leer ist.

    Ablauf je Eintrag:
      1. Lese Eigentümer, Adresse, PLZ/Ort, Bundesland aus Notion.
      2. Bestimme zuständige Person aus KONTAKT_DATEN.
      3. Erzeuge Anrede (geschlechtsspezifisch).
      4. Befülle DOCX-Vorlage (brief_vorlage.docx).
      5. Sende DOCX per E-Mail an Betreuer (Option C).
      6. Speichere DOCX lokal als GitHub-Artifact (Backup).
      7. Setze 'Brief erstellt am' in Notion.
      8. Füge Notiz "Brief erstellt am DD.MM.YYYY" hinzu.

    Gibt (Anzahl erstellter Briefe, Liste der Telegram-Zeilen) zurück.
    """
    if not DOCX_AVAILABLE:
        print("[Brief] ⚠️  python-docx nicht installiert – überspringe Brief-Erstellung")
        return 0, []

    if not os.path.exists(BRIEF_VORLAGE_PATH):
        print(f"[Brief] ⚠️  Vorlage nicht gefunden: {BRIEF_VORLAGE_PATH} – überspringe")
        return 0, []

    ZIEL_PHASE = "✅ Relevant – Brief vorbereiten"

    print("\n[Brief] 📝 Suche nach Einträgen für Brief-Erstellung …")

    pages = all_pages if all_pages is not None else notion_load_all_pages(notion, db_id)

    to_process: list[dict] = []
    for page in pages:
        props = page.get("properties", {})
        phase = (props.get("Workflow-Phase", {}).get("select") or {}).get("name", "")
        if phase != ZIEL_PHASE:
            continue
        # Überspringe wenn Brief bereits erstellt (per Datumsfeld ODER Notiz-Marker)
        brief_datum = props.get("Brief erstellt am", {}).get("date")
        if brief_datum and brief_datum.get("start"):
            continue
        # Fallback: prüfe ob Notiz bereits "Brief erstellt am" enthält
        notizen_rt = props.get("Notizen", {}).get("rich_text", [])
        notizen_text = "".join(t.get("plain_text", "") for t in notizen_rt)
        if "Brief erstellt am" in notizen_text:
            continue
        to_process.append(page)

    print(f"[Brief] 📋 {len(to_process)} Einträge für Brief-Erstellung gefunden")
    if not to_process:
        return 0, []

    erstellt = 0
    telegram_lines: list[str] = []
    from datetime import date

    # Ausgabe-Verzeichnis für DOCXs (wird als GitHub-Artifact hochgeladen)
    brief_output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "briefe")
    os.makedirs(brief_output_dir, exist_ok=True)

    # ── Einträge nach Eigentümer gruppieren ───────────────────────────────────
    # Gleicher Eigentümer → ein Brief mit allen Liegenschaften aufgelistet
    from collections import defaultdict
    gruppen: dict[str, list[dict]] = defaultdict(list)

    for page in to_process:
        props = page.get("properties", {})
        eigentuemer_list = props.get("Verpflichtende Partei", {}).get("rich_text", [])
        eigentuemer      = "".join(t.get("plain_text", "") for t in eigentuemer_list).strip()
        # Normalisierter Key: Kleinbuchstaben, Leerzeichen zusammengefasst
        key = " ".join(eigentuemer.lower().split()) if eigentuemer else f"__leer_{page['id']}"
        gruppen[key].append(page)

    print(f"[Brief] 👥 {len(gruppen)} Eigentümer-Gruppe(n) → {len(to_process)} Einträge")

    heute     = date.today()
    datum_str = heute.strftime("%d.%m.%Y")

    for eigen_key, gruppe in gruppen.items():
        # ── Daten aus erstem Eintrag der Gruppe lesen ─────────────────────────
        first_page  = gruppe[0]
        first_props = first_page.get("properties", {})

        eigentuemer_list = first_props.get("Verpflichtende Partei", {}).get("rich_text", [])
        eigentuemer      = "".join(t.get("plain_text", "") for t in eigentuemer_list).strip()

        adresse_list = first_props.get("Zustell Adresse", {}).get("rich_text", [])
        adresse      = "".join(t.get("plain_text", "") for t in adresse_list).strip()

        plz_ort_list = first_props.get("Zustell PLZ/Ort", {}).get("rich_text", [])
        plz_ort      = "".join(t.get("plain_text", "") for t in plz_ort_list).strip()

        bundesland = (first_props.get("Bundesland", {}).get("select") or {}).get("name", "")

        # ── Pflichtfelder prüfen ──────────────────────────────────────────────
        titel_list = first_props.get("Liegenschaftsadresse", {}).get("title", [])
        titel      = "".join(t.get("plain_text", "") for t in titel_list).strip()

        if not eigentuemer:
            print(f"  [Brief] ⏭  Überspringe Gruppe – kein Eigentümer")
            continue
        if not adresse or not plz_ort:
            print(f"  [Brief] ⏭  Überspringe {eigentuemer[:50]} – keine Zustelladresse")
            continue
        if not bundesland or bundesland not in KONTAKT_DATEN:
            print(f"  [Brief] ⏭  Überspringe {eigentuemer[:50]} – kein Kontakt für '{bundesland}'")
            continue

        kontakt = KONTAKT_DATEN[bundesland]

        # ── Alle Liegenschaften der Gruppe sammeln ────────────────────────────
        liegenschaften = []
        for page in gruppe:
            props = page.get("properties", {})

            # ── Liegenschaftsadresse: bevorzugt das dedizierte Titelfeld ──────
            # Notion speichert die volle Adresse (Straße + Hausnr) im Titelfeld
            # "Liegenschaftsadresse", PLZ/Ort separat in "Liegenschafts PLZ"
            titel_rt  = props.get("Liegenschaftsadresse", {}).get("title", [])
            t_adresse = "".join(x.get("plain_text", "") for x in titel_rt).strip()

            # Fallback: altes "Name"-Feld
            if not t_adresse:
                t_list    = props.get("Name", {}).get("title", [])
                t_adresse = "".join(x.get("plain_text", "") for x in t_list).strip()

            # PLZ/Ort: zuerst dediziertes Feld, dann aus Adresse extrahieren
            plz_rt    = props.get(NOTION_PLZ_FIELD, {}).get("rich_text", [])
            t_plz_ort = "".join(x.get("plain_text", "") for x in plz_rt).strip()

            if not t_plz_ort:
                # Versuche PLZ/Ort am Ende der Adresse zu finden (z.B. "Musterstr. 1, 1010 Wien")
                t_parts = t_adresse.rsplit(",", 1)
                if len(t_parts) == 2 and re.match(r"\s*\d{4}", t_parts[1]):
                    t_adresse = t_parts[0].strip()
                    t_plz_ort = t_parts[1].strip()

            liegenschaften.append({
                "adresse": t_adresse,
                "plz_ort": t_plz_ort,
                "titel":   t_adresse,
            })

        # ── Liegenschaft(en) für Platzhalter aufbereiten ──────────────────────
        # Bei mehreren: erste Liegenschaft im Template, Rest als Auflistung
        erste = liegenschaften[0]
        if len(liegenschaften) == 1:
            liegenschaft_adresse = erste["adresse"]
            liegenschaft_plz_ort = erste["plz_ort"] or plz_ort
        else:
            # Mehrere Liegenschaften → im Betreff aufzählen
            liegenschaft_adresse = erste["adresse"]
            liegenschaft_plz_ort = erste["plz_ort"] or plz_ort
            zusatz = "\n".join(
                f"  • {l['adresse']}{', ' + l['plz_ort'] if l['plz_ort'] else ''}"
                for l in liegenschaften[1:]
            )
            liegenschaft_plz_ort += f"\n\nWeitere Liegenschaften:\n{zusatz}"

        # ── Anrede + Platzhalter ──────────────────────────────────────────────
        anrede = _brief_anrede(eigentuemer)

        platzhalter = {
            "EIGENTUEMER_NAME":     eigentuemer,
            "ZUSTELL_ADRESSE":      adresse,
            "ZUSTELL_PLZ_ORT":      plz_ort,
            "DATUM":                f"Wien, am {datum_str}",
            "LIEGENSCHAFT_ADRESSE": liegenschaft_adresse,
            "LIEGENSCHAFT_PLZ_ORT": liegenschaft_plz_ort,
            "ANREDE":               anrede,
            "KONTAKT_NAME":         kontakt["name"],
            "KONTAKT_TEL":          kontakt["tel"],
            "KONTAKT_EMAIL":        kontakt["email"],
        }

        try:
            # ── DOCX befüllen ─────────────────────────────────────────────────
            docx_bytes = _brief_fill_template(BRIEF_VORLAGE_PATH, platzhalter)

            # ── Dateiname ─────────────────────────────────────────────────────
            safe_eigen     = re.sub(r"[^\w\s-]", "", eigentuemer)[:40].strip().replace(" ", "_")
            safe_datum     = datum_str.replace(".", "-")
            dateiname_docx = f"Brief_{safe_datum}_{safe_eigen}.docx"
            docx_path      = os.path.join(brief_output_dir, dateiname_docx)

            with open(docx_path, "wb") as f:
                f.write(docx_bytes)

            anzahl_str = f" ({len(liegenschaften)} Liegenschaften)" if len(liegenschaften) > 1 else ""
            print(f"  [Brief] 💾 DOCX gespeichert: {dateiname_docx}{anzahl_str}")

            # ── E-Mail an Betreuer ────────────────────────────────────────────
            email_ok = _brief_send_email(
                kontakt_email  = kontakt["email"],
                kontakt_name   = kontakt["name"],
                eigentuemer    = eigentuemer,
                titel          = titel if len(liegenschaften) == 1 else
                                 f"{len(liegenschaften)} Liegenschaften: " +
                                 ", ".join(l["adresse"][:30] for l in liegenschaften[:3]),
                docx_bytes     = docx_bytes,
                dateiname_docx = dateiname_docx,
            )

            # ── Telegram-Dokument ─────────────────────────────────────────────
            tg_caption = (
                f"📄 Brief für {eigentuemer[:60]}\n"
                f"📍 {bundesland} | Betreuer: {kontakt['name']}\n"
                f"📅 {datum_str}"
                + (f"\n🏠 {len(liegenschaften)} Liegenschaften" if len(liegenschaften) > 1 else "")
            )
            send_telegram_document(docx_bytes, dateiname_docx, caption=tg_caption, bundesland=bundesland)

            # ── Notion: alle Seiten der Gruppe aktualisieren ──────────────────
            versand_info = f"E-Mail an {kontakt['email']}" if email_ok else "Telegram"
            for page in gruppe:
                p_id = page["id"]
                p_props = page.get("properties", {})
                notizen_list = p_props.get("Notizen", {}).get("rich_text", [])
                notizen_alt  = "".join(t.get("plain_text", "") for t in notizen_list).strip()
                neue_notiz   = notizen_alt
                if neue_notiz and not neue_notiz.endswith("\n"):
                    neue_notiz += "\n"
                if len(liegenschaften) > 1:
                    neue_notiz += f"Brief erstellt am {datum_str} (Sammelbrief, {versand_info})"
                else:
                    neue_notiz += f"Brief erstellt am {datum_str} ({versand_info})"
                neue_notiz = neue_notiz[:2000]

                try:
                    notion.pages.update(
                        page_id=p_id,
                        properties={
                            "Brief erstellt am": {"date": {"start": heute.isoformat()}},
                            "Notizen": {"rich_text": [{"type": "text", "text": {"content": neue_notiz}}]},
                        }
                    )
                except Exception as notion_exc:
                    err_str = str(notion_exc)
                    if "Brief erstellt am" in err_str and "not a property" in err_str:
                        try:
                            notion.pages.update(
                                page_id=p_id,
                                properties={
                                    "Notizen": {"rich_text": [{"type": "text", "text": {"content": neue_notiz}}]},
                                }
                            )
                        except Exception as notiz_exc:
                            print(f"  [Brief] ⚠️  Notiz-Update fehlgeschlagen für {p_id[:8]}: {notiz_exc}")
                    else:
                        print(f"  [Brief] ⚠️  Notion-Update fehlgeschlagen für {p_id[:8]}: {notion_exc}")
                time.sleep(0.2)

            icon = "✉️" if email_ok else "📨"
            print(f"  [Brief] ✅ Erledigt: {eigentuemer[:40]} ({bundesland}) → {kontakt['name']}{anzahl_str}")
            telegram_lines.append(
                f"{icon} {html_escape(eigentuemer[:35])} | {html_escape(bundesland)} "
                f"→ {html_escape(kontakt['name'])}{html_escape(anzahl_str)}"
            )
            erstellt += 1
            time.sleep(0.3)

        except Exception as exc:
            print(f"  [Brief] ❌ Fehler bei {eigentuemer[:50]}: {exc}")

    print(f"[Brief] ✅ {erstellt} Brief(e) erstellt ({len(to_process)} Einträge)")
    return erstellt, telegram_lines


def fetch_results_for_state(bundesland: str, bl_value: str) -> list[dict]:
    """
    Ruft die Ergebnisseite für ein Bundesland direkt per HTTP ab.

    Die URL-Struktur wurde durch Analyse des Formulars ermittelt:
    /edikte/ex/exedi3.nsf/suchedi?SearchView&subf=eex&...&query=([BL]=(X))
    """
    print(f"\n[Scraper] 🔍 Suche für: {bundesland} (BL={bl_value})")

    query = f"([BL]=({bl_value}))"
    params = urllib.parse.urlencode({
        "SearchView": "",
        "subf": "eex",
        "SearchOrder": "4",
        "SearchMax": "4999",
        "retfields": f"~BL={bl_value}",
        "ftquery": "",
        "query": query,
    })
    url = f"{BASE_URL}/edikte/ex/exedi3.nsf/suchedi?{params}"

    try:
        req = urllib.request.Request(
            url,
            headers={"User-Agent": "Mozilla/5.0 (compatible; EdikteMonitor/1.0)"}
        )
        with urllib.request.urlopen(req, timeout=30) as r:
            html = r.read().decode("utf-8", errors="replace")
    except Exception as exc:
        print(f"  [Scraper] ❌ HTTP-Fehler: {exc}")
        return []

    # Links extrahieren – Format: alldoc/HEX!OpenDocument (relativ, ohne führendes /)
    rel_pattern = re.compile(
        r'<a[^>]+href="(alldoc/([0-9a-f]+)!OpenDocument)"[^>]*>([^<]+)</a>',
        re.IGNORECASE
    )

    results = []
    seen_ids = set()

    for href_rel, edikt_id, link_text in rel_pattern.findall(html):
        link_text = link_text.strip()
        edikt_id  = edikt_id.lower()
        href      = f"{BASE_URL}/edikte/ex/exedi3.nsf/{href_rel}"

        if edikt_id in seen_ids:
            continue
        seen_ids.add(edikt_id)

        # Typ bestimmen
        typ = None
        for t in RELEVANT_TYPES:
            if link_text.startswith(t):
                typ = t
                break
        if not typ:
            continue

        # Ausschlussliste (nur bei Versteigerung relevant)
        if typ == "Versteigerung" and is_excluded(link_text):
            print(f"  [Filter] ⛔ Ausgeschlossen: {link_text[:80]}")
            continue

        results.append({
            "bundesland":   bundesland,
            "type":         typ,
            "beschreibung": link_text,
            "link":         href,
            "edikt_id":     edikt_id,
        })

    print(f"  [Scraper] 📋 {len(results)} relevante Treffer für {bundesland}")
    return results


# =============================================================================
# MAIN
# =============================================================================

async def main() -> None:
    print("=" * 60)
    print(f"Edikte-Monitor gestartet: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 60)

    notion = Client(auth=env("NOTION_TOKEN"))
    db_id  = clean_notion_db_id(env("NOTION_DATABASE_ID"))

    # ── BRIEF_ONLY-Modus: nur Status-Sync + Brief-Erstellung ─────────────────
    # Wird gesetzt wenn Env-Variable BRIEF_ONLY=true gesetzt ist.
    # Kein Scraping, keine PDF-Analyse – läuft in ~30 Sekunden statt ~10 Minuten.
    if os.environ.get("BRIEF_ONLY", "").lower() == "true":
        print("[Modus] ⚡ BRIEF_ONLY – nur Status-Sync + Brief-Erstellung")
        try:
            _pages = notion_load_all_pages(notion, db_id)
            notion_status_sync(notion, db_id, all_pages=_pages)
            # Seiten nach Sync neu laden damit aktualisierte Phasen sichtbar sind
            _pages = notion_load_all_pages(notion, db_id)

            # ── Google Drive: Unterlagen für alle Gelb-Einträge hochladen ────
            _gdrive_service = gdrive_get_service()
            if _gdrive_service:
                gdrive_clear_placeholder_links(notion, db_id, _pages)
                gdrive_sync_gelb_entries(notion, db_id, _pages, _gdrive_service)
            else:
                print("[GDrive] ℹ️  Kein Service verfügbar (Bibliothek nicht installiert oder Key fehlt)")

            brief_erstellt, brief_telegram = notion_brief_erstellen(notion, db_id, all_pages=_pages)
            print(f"[Modus] ✅ BRIEF_ONLY fertig – {brief_erstellt} Brief(e) erstellt")
            if brief_erstellt == 0:
                print("[Modus] Keine neuen Briefe – kein Telegram-Versand.")
            else:
                lines = [
                    "<b>📨 Neue Briefe erstellt</b>",
                    f"<i>{datetime.now().strftime('%d.%m.%Y %H:%M')}</i>",
                    "",
                    f"<b>✉️ Briefe erstellt: {brief_erstellt}</b>",
                ]
                for bl in brief_telegram[:20]:
                    lines.append(f"• {bl}")
                await send_telegram("\n".join(lines))
        except Exception as exc:
            print(f"[Modus] ❌ BRIEF_ONLY Fehler: {exc}")
        return

    neue_eintraege:  list[dict] = []
    edikt_updates:   list[str]  = []   # Telegram-Zeilen für Edikt-Updates
    fehler:          list[str]  = []

    # ── 1. Alle bekannten IDs einmalig laden (schnelle lokale Deduplizierung) ─
    try:
        known_ids = notion_load_all_ids(notion, db_id)  # {edikt_id -> page_id}
    except Exception as exc:
        print(f"  [ERROR] Konnte IDs nicht laden: {exc}")
        known_ids = {}

    # ── 2. Edikte scrapen + in Notion eintragen ───────────────────────────────
    for bundesland, bl_value in BUNDESLAENDER.items():
        try:
            results = fetch_results_for_state(bundesland, bl_value)
        except Exception as exc:
            msg = f"Scraper-Fehler {bundesland}: {exc}"
            print(f"  [ERROR] {msg}")
            fehler.append(msg)
            continue

        for item in results:
            try:
                eid = item["edikt_id"].lower()

                if item["type"] == "Versteigerung":
                    if known_ids.get(eid) == "(geschuetzt)":
                        print(f"  [Notion] 🔒 Geschützt (bereits bearbeitet): {eid}")
                    elif eid not in known_ids:
                        result_tuple = notion_create_eintrag(notion, db_id, item, known_ids=known_ids)
                        if result_tuple is None:
                            # Kategorie-Filter oder Titel-Duplikat ohne Update-Info
                            known_ids[eid] = "(gefiltert)"
                        elif isinstance(result_tuple, tuple) and result_tuple[0] == "__edikt_update__":
                            # Selbe Immobilie, neue edikt_id → bestehenden Eintrag updaten
                            _, existing_page_id, detail = result_tuple
                            notion_update_edikt_eintrag(notion, existing_page_id, item, detail)
                            known_ids[eid] = "(geschuetzt)"
                            titel_rt = detail.get("adresse_voll") or item.get("beschreibung", "")[:60]
                            termin   = detail.get("termin_iso", "")
                            edikt_updates.append(
                                f"🔄 <b>{titel_rt[:70]}</b>"
                                + (f"\nNeuer Termin: {termin}" if termin else "")
                            )
                        else:
                            detail, new_page_id = result_tuple
                            item["_detail"] = detail
                            neue_eintraege.append(item)
                            known_ids[eid] = "(neu)"  # sofort als bekannt markieren
                            # ── Gutachten sofort anreichern ──────────────────
                            if new_page_id and item.get("link") and FITZ_AVAILABLE:
                                try:
                                    gutachten_enrich_notion_page(
                                        notion, new_page_id, item["link"]
                                    )
                                except Exception as ge:
                                    print(f"    [Gutachten] ⚠️  Anreicherung fehlgeschlagen: {ge}")
                    else:
                        print(f"  [Notion] ⏭  Bereits vorhanden: {eid}")

                elif item["type"] in ("Entfall des Termins", "Verschiebung"):
                    page_id = known_ids.get(eid)
                    if page_id and page_id not in ("(neu)", "(geschuetzt)", "(gefiltert)"):
                        notion_mark_entfall(notion, page_id, item)
                        # Kein Telegram für Entfall/Verschiebung – nur Notion-Eintrag
                    elif page_id == "(geschuetzt)":
                        print(f"  [Notion] 🔒 Entfall übersprungen (geschützte Phase): {eid}")
                    else:
                        # Kein DB-Eintrag → nur loggen, keine Benachrichtigung
                        print(f"  [Notion] ℹ️  Entfall ohne DB-Eintrag: {eid}")

            except Exception as exc:
                msg = f"Notion-Fehler {item.get('edikt_id', '?')}: {exc}"
                print(f"  [ERROR] {msg}")
                fehler.append(msg)

    # ── 3. URL-Anreicherung für manuell angelegte Einträge ────────────────────
    try:
        enriched_count = notion_enrich_urls(notion, db_id)
    except Exception as exc:
        msg = f"URL-Anreicherung fehlgeschlagen: {exc}"
        print(f"  [ERROR] {msg}")
        fehler.append(msg)
        enriched_count = 0

    # ── 3. Einmaliges Laden aller Notion-Pages ────────────────────────────────
    # Die folgenden 4 Schritte (Status-Sync, Bereinigung, Tote-URLs,
    # Qualitäts-Check) würden sonst jeweils einen eigenen DB-Scan starten.
    # Stattdessen laden wir die DB EINMALIG und geben das Ergebnis weiter.
    try:
        _all_pages = notion_load_all_pages(notion, db_id)
    except Exception as exc:
        print(f"  [WARN] Konnte Pages nicht vorladen – Fallback auf Einzel-Scans: {exc}")
        _all_pages = None   # jede Funktion macht dann selbst einen Scan

    # ── 3a. Status-Sync: Status-Farbe / Für-uns-relevant? → Phase + Checkboxen ─
    # Wenn ein Kollege manuell 🔴/🟡/🟢 setzt oder "Für uns relevant?" befüllt,
    # werden Phase und Checkboxen automatisch angepasst (kein manuelles Ankreuzen nötig).
    try:
        notion_status_sync(notion, db_id, all_pages=_all_pages)
    except Exception as exc:
        print(f"  [WARN] Status-Sync fehlgeschlagen (nicht kritisch): {exc}")

    # ── WICHTIG: Pages nach Status-Sync neu laden ────────────────────────────
    # Status-Sync hat Phasen/Checkboxen in Notion aktualisiert.
    # Damit Brief-Erstellung und Qualitäts-Check die neuen Werte sehen,
    # muss die lokale Kopie jetzt neu geladen werden.
    try:
        _all_pages = notion_load_all_pages(notion, db_id)
    except Exception as exc:
        print(f"  [WARN] Neu-Laden nach Status-Sync fehlgeschlagen – Fallback auf alte Daten: {exc}")

    # ── 3b. Einmalige Bereinigung: falsche Gerichtsnamen in 'Verpflichtende Partei' ──
    # Frühere Script-Versionen haben irrtümlich den Gerichtsnamen (z.B. "BG Schwaz (870)")
    # in das Feld 'Verpflichtende Partei' geschrieben. Diese Einträge werden hier
    # erkannt, das Feld geleert und 'Gutachten analysiert?' zurückgesetzt,
    # damit der nächste Schritt (4) sie neu verarbeitet.
    try:
        notion_reset_falsche_verpflichtende(notion, db_id, all_pages=_all_pages)
    except Exception as exc:
        print(f"  [WARN] Bereinigung fehlgeschlagen (nicht kritisch): {exc}")

    # ── 3c. Tote URLs archivieren (HTTP 404) ─────────────────────────────────
    tote_urls_archiviert = 0
    tote_urls_alarme: list[str] = []
    try:
        tote_urls_archiviert, tote_urls_alarme = notion_archiviere_tote_urls(notion, db_id, all_pages=_all_pages)
    except Exception as exc:
        print(f"  [WARN] Tote-URLs-Check fehlgeschlagen (nicht kritisch): {exc}")

    # ── 3d. Qualitäts-Check: analysierte Einträge auf Vollständigkeit prüfen ──
    # Einträge die als 'analysiert' markiert sind, aber keinen Eigentümer/Adresse
    # haben, werden zurückgesetzt damit Schritt 4 sie neu analysiert.
    try:
        notion_qualitaetscheck(notion, db_id, all_pages=_all_pages)
    except Exception as exc:
        print(f"  [WARN] Qualitäts-Check fehlgeschlagen (nicht kritisch): {exc}")

    # ── 3e. Google Drive: Unterlagen für Gelb-Einträge hochladen ────────────────
    try:
        _gdrive_service = gdrive_get_service()
        if _gdrive_service:
            gdrive_clear_placeholder_links(notion, db_id, _all_pages or [])
            gdrive_sync_gelb_entries(notion, db_id, _all_pages or [], _gdrive_service)
        else:
            print("[GDrive] ℹ️  Kein Service verfügbar (Bibliothek nicht installiert oder Key fehlt)")
    except Exception as exc:
        print(f"  [WARN] Google Drive Sync fehlgeschlagen (nicht kritisch): {exc}")

    # ── 3f. Brief-Erstellung: relevant markierte Einträge → Brief erstellen ──────────────
    # Betrifft: Einträge mit Phase '✅ Relevant – Brief vorbereiten'
    # bei denen 'Brief erstellt am' noch leer ist.
    brief_erstellt = 0
    brief_telegram: list[str] = []
    try:
        brief_erstellt, brief_telegram = notion_brief_erstellen(notion, db_id, all_pages=_all_pages)
    except Exception as exc:
        print(f"  [WARN] Brief-Erstellung fehlgeschlagen (nicht kritisch): {exc}")

    # ── 4. Gutachten-Anreicherung: Text-PDFs (LLM) ───────────────────────────
    # Betrifft: Einträge die eine URL haben aber noch nicht analysiert wurden.
    gutachten_enriched = 0
    if FITZ_AVAILABLE:
        try:
            gutachten_enriched = notion_enrich_gutachten(notion, db_id)
        except Exception as exc:
            msg = f"Gutachten-Anreicherung fehlgeschlagen: {exc}"
            print(f"  [ERROR] {msg}")
            fehler.append(msg)
    else:
        print("[Gutachten] ℹ️  PyMuPDF nicht verfügbar – überspringe Gutachten-Anreicherung")

    # ── 4b. Vision-Analyse: gescannte PDFs (GPT-4o) ──────────────────────────
    vision_enriched = 0
    try:
        vision_enriched = notion_enrich_gescannte(notion, db_id)
    except Exception as exc:
        print(f"  [WARN] Vision-Analyse fehlgeschlagen (nicht kritisch): {exc}")

    # ── 5. Zusammenfassung ────────────────────────────────────────────────────
    print("\n" + "=" * 60)
    print(f"✅ Neue Einträge:         {len(neue_eintraege)}")
    print(f"🔗 URLs ergänzt:          {enriched_count}")
    print(f"🗑  Tote URLs archiviert:  {tote_urls_archiviert}")
    print(f"📄 Gutachten analysiert:  {gutachten_enriched}")
    print(f"🔭 Vision analysiert:     {vision_enriched}")
    print(f"✉️  Briefe erstellt:      {brief_erstellt}")
    print(f"⚠️  Fehler:                {len(fehler)}")
    print("=" * 60)

    if not neue_eintraege and not fehler \
            and not gutachten_enriched and not vision_enriched \
            and not tote_urls_archiviert and not tote_urls_alarme \
            and not brief_erstellt and not edikt_updates:
        print("Keine neuen relevanten Änderungen – kein Telegram-Versand.")
        return

    # ── 6. Telegram ───────────────────────────────────────────────────────────
    lines = [
        "<b>🏛 Edikte-Monitor</b>",
        f"<i>{datetime.now().strftime('%d.%m.%Y %H:%M')}</i>",
        "",
    ]

    if neue_eintraege:
        lines.append(f"<b>🟢 Neue Versteigerungen: {len(neue_eintraege)}</b>")
        for item in neue_eintraege[:20]:
            detail    = item.get("_detail", {})
            adresse   = html_escape(detail.get("adresse_voll") or item["beschreibung"][:70])
            kategorie = html_escape(detail.get("kategorie", ""))
            vk        = detail.get("schaetzwert")
            vk_str    = f" | 💰 {vk:,.0f} €".replace(",", ".") if vk else ""
            kat_str   = f" [{kategorie}]" if kategorie else ""
            lines.append(f"• <b>{adresse}</b>{kat_str}{vk_str}")
            lines.append(f"  <a href=\"{item['link']}\">→ Edikt öffnen</a>")
        if len(neue_eintraege) > 20:
            lines.append(f"  ... und {len(neue_eintraege) - 20} weitere")
        lines.append("")

    # Entfall/Verschiebung wird nicht per Telegram gemeldet – nur in Notion eingetragen

    if enriched_count:
        lines.append(f"<b>🔗 URLs nachgetragen: {enriched_count}</b>")
        lines.append("")

    if tote_urls_archiviert:
        lines.append(f"<b>🗑 Tote Edikte archiviert: {tote_urls_archiviert}</b>")
        lines.append("")

    if edikt_updates:
        lines.append(f"<b>🔄 Edikt-Updates (gleiche Immobilie, neuer Termin): {len(edikt_updates)}</b>")
        for upd in edikt_updates[:10]:
            lines.append(f"• {upd}")
        lines.append("")

    if tote_urls_alarme:
        lines.append("<b>🚨 Achtung – Edikt verschwunden (manuelle Prüfung!):</b>")
        for alarm in tote_urls_alarme:
            lines.append(f"• {alarm}")
        lines.append("")

    if brief_erstellt:
        lines.append(f"<b>✉️ Briefe erstellt: {brief_erstellt}</b>")
        for bl in brief_telegram[:10]:
            lines.append(f"• {bl}")
        lines.append("")

    if gutachten_enriched:
        lines.append(f"<b>📄 Gutachten analysiert (Text): {gutachten_enriched}</b>")
        lines.append("")

    if vision_enriched:
        lines.append(f"<b>🔭 Gutachten analysiert (Vision): {vision_enriched}</b>")
        lines.append("")

    if fehler:
        lines.append(f"<b>⚠️ Fehler ({len(fehler)}):</b>")
        for f_msg in fehler[:5]:
            lines.append(f"• {f_msg[:100]}")

    try:
        await send_telegram("\n".join(lines))
    except Exception as exc:
        print(f"[ERROR] Telegram fehlgeschlagen: {exc}")

    # ── Gefilterte Nachrichten direkt an Betreuer senden ─────────────────────
    tg_token = env("TELEGRAM_BOT_TOKEN")
    tg_url   = f"https://api.telegram.org/bot{tg_token}/sendMessage"

    def _send_filtered(chat_id: str, name: str, bundeslaender: set, label: str) -> None:
        """Sendet gefilterte Versteigerungs-Nachricht direkt an einen Betreuer."""
        eintraege = [
            e for e in neue_eintraege
            if e.get("bundesland", "") in bundeslaender
        ]
        if not eintraege:
            print(f"[Telegram] ℹ️  Keine {label}-Einträge – kein Telegram an {name}")
            return
        lines = [
            "<b>🏛 Edikte-Monitor</b>",
            f"<i>{datetime.now().strftime('%d.%m.%Y %H:%M')}</i>",
            f"<i>({html_escape(label)})</i>",
            "",
            f"<b>🟢 Neue Versteigerungen: {len(eintraege)}</b>",
        ]
        for item in eintraege[:20]:
            detail    = item.get("_detail", {})
            adresse   = html_escape(detail.get("adresse_voll") or item["beschreibung"][:70])
            kategorie = html_escape(detail.get("kategorie", ""))
            kat_str   = f" [{kategorie}]" if kategorie else ""
            lines.append(f"• {html_escape(item['bundesland'])} – {adresse}{kat_str}")
        msg = "\n".join(lines)
        try:
            _telegram_send_raw(tg_url, {
                "chat_id":                  chat_id,
                "text":                     msg,
                "parse_mode":               "HTML",
                "disable_web_page_preview": True,
            })
            print(f"[Telegram] ✅ Gefilterte Nachricht an {name} gesendet ({len(eintraege)} Einträge)")
        except Exception:
            plain = _truncate_plain(_strip_html_tags(msg))
            try:
                _telegram_send_raw(tg_url, {
                    "chat_id":                  chat_id,
                    "text":                     plain,
                    "disable_web_page_preview": True,
                })
                print(f"[Telegram] ✅ Gefilterte Nachricht an {name} (Plain) gesendet ({len(eintraege)} Einträge)")
            except Exception as exc2:
                print(f"[ERROR] Telegram {name} fehlgeschlagen: {exc2}")

    # Benjamin: Wien + Oberösterreich
    benjamin_id = _get_benjamin_chat_id()
    if benjamin_id:
        try:
            _send_filtered(benjamin_id, "Benjamin", BENJAMIN_BUNDESLAENDER, "Wien & Oberösterreich")
        except Exception as exc:
            print(f"[ERROR] Telegram Benjamin fehlgeschlagen: {exc}")

    # Christopher: Niederösterreich + Burgenland
    christopher_id = _get_christopher_chat_id()
    if christopher_id:
        try:
            _send_filtered(christopher_id, "Christopher", CHRISTOPHER_BUNDESLAENDER, "Niederösterreich & Burgenland")
        except Exception as exc:
            print(f"[ERROR] Telegram Christopher fehlgeschlagen: {exc}")


if __name__ == "__main__":
    asyncio.run(main())
